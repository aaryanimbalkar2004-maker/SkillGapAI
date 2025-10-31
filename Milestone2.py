"""
COMPLETE Milestone 2: AI-Powered Skill Gap Analyzer - Skill Extraction
Integrated implementation with ALL features:
- Multi-method skill extraction (spaCy, POS, context, NER)
- BERT embeddings with Sentence-BERT
- Custom NER training capability
- Annotation interface for training data
- Comprehensive visualizations
- Multiple export formats

Requirements:
pip install streamlit spacy scikit-learn pandas numpy plotly sentence-transformers

python -m spacy download en_core_web_sm

Run with: streamlit run milestone2.py
"""

import streamlit as st
import spacy
from spacy.training import Example
import pandas as pd
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.graph_objects as go
import plotly.express as px
import re
import json

import streamlit as st
import PyPDF2
import docx
import re

import random
import logging
from typing import Dict, List, Set, Tuple, Optional
from collections import Counter, defaultdict
from datetime import datetime
from io import BytesIO
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import spacy
from spacy.training import Example
from sentence_transformers import SentenceTransformer
from collections import defaultdict


# Configure page
st.set_page_config(
    page_title="AI Skill Gap Analyzer - Milestone 2",
    page_icon="🎯",
    layout="wide"
)

# Custom CSS
st.markdown("""
    <style>
    .skill-tag {
        display: inline-block;
        padding: 5px 10px;
        margin: 5px;
        background-color: skyblue;
        border-radius: 15px;
        font-size: 14px;
    }
    .tech-skill {
        background-color: #e3f2fd;
        color: #1976d2;
    }
    .soft-skill {
        background-color: #f3e5f5;
        color: #7b1fa2;
    }
    </style>
""", unsafe_allow_html=True)




class SkillDatabase:
    """Comprehensive skill database with categorization"""
    
    def __init__(self):
        self.skills = self._initialize_skill_database()
        self.abbreviations = self._initialize_abbreviations()
        self.skill_patterns = self._initialize_skill_patterns()
    
    def _initialize_skill_database(self) -> Dict[str, List[str]]:
        """Initialize comprehensive skill database"""
        return {
            'programming_languages': [
                'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C',
                'Ruby', 'PHP', 'Swift', 'Kotlin', 'Go', 'Rust', 'Scala', 'R',
                'MATLAB', 'Perl', 'Dart', 'Shell', 'Bash', 'PowerShell'
            ],
            'web_frameworks': [
                'React', 'Angular', 'Vue.js', 'Node.js', 'Express.js', 'Django',
                'Flask', 'FastAPI', 'Spring Boot', 'Spring', 'ASP.NET', '.NET Core',
                'Ruby on Rails', 'Laravel', 'Next.js', 'Nuxt.js', 'Svelte',
                'jQuery', 'Bootstrap', 'Tailwind CSS', 'Material-UI'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Cassandra', 'Oracle',
                'SQL Server', 'SQLite', 'MariaDB', 'DynamoDB', 'Elasticsearch',
                'Firebase', 'Neo4j', 'Snowflake', 'BigQuery'
            ],
            'ml_ai': [
                'Machine Learning', 'Deep Learning', 'Neural Networks',
                'Natural Language Processing', 'NLP', 'Computer Vision',
                'Reinforcement Learning', 'Transfer Learning',
                'Feature Engineering', 'MLOps', 'Generative AI',
                'Large Language Models', 'LLM', 'CNN', 'RNN', 'LSTM',
                'Transformer', 'BERT', 'GPT'
            ],
            'ml_frameworks': [
                'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'XGBoost',
                'Pandas', 'NumPy', 'SciPy', 'Matplotlib', 'Seaborn',
                'Plotly', 'NLTK', 'spaCy', 'Hugging Face', 'OpenCV'
            ],
            'cloud_platforms': [
                'AWS', 'Amazon Web Services', 'Azure', 'Microsoft Azure',
                'Google Cloud Platform', 'GCP', 'Heroku', 'DigitalOcean'
            ],
            'devops_tools': [
                'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI', 'GitHub Actions',
                'CircleCI', 'Ansible', 'Terraform', 'Prometheus', 'Grafana',
                'ELK Stack', 'Datadog'
            ],
            'version_control': [
                'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN'
            ],
            'testing': [
                'Jest', 'Mocha', 'Pytest', 'JUnit', 'Selenium', 'Cypress',
                'Postman', 'JMeter'
            ],
            'soft_skills': [
                'Leadership', 'Team Management', 'Communication',
                'Problem Solving', 'Critical Thinking', 'Analytical Skills',
                'Project Management', 'Collaboration', 'Teamwork',
                'Adaptability', 'Creativity', 'Time Management'
            ]
        }
    
    def _initialize_abbreviations(self) -> Dict[str, str]:
        """Initialize common abbreviations"""
        return {
            'ML': 'Machine Learning', 'DL': 'Deep Learning',
            'AI': 'Artificial Intelligence', 'NLP': 'Natural Language Processing',
            'CV': 'Computer Vision', 'NN': 'Neural Networks',
            'CNN': 'Convolutional Neural Networks', 'RNN': 'Recurrent Neural Networks',
            'K8s': 'Kubernetes', 'K8S': 'Kubernetes',
            'CI/CD': 'Continuous Integration/Continuous Deployment',
            'API': 'Application Programming Interface',
            'REST': 'Representational State Transfer',
            'SQL': 'Structured Query Language', 'OOP': 'Object-Oriented Programming',
            'TDD': 'Test-Driven Development', 'AWS': 'Amazon Web Services',
            'GCP': 'Google Cloud Platform'
        }
    
    def _initialize_skill_patterns(self) -> List[str]:
        """Initialize regex patterns for skill detection"""
        return [
            r'experience (?:in|with) ([\w\s\+\#\.\-]+)',
            r'proficient (?:in|with|at) ([\w\s\+\#\.\-]+)',
            r'expertise (?:in|with) ([\w\s\+\#\.\-]+)',
            r'knowledge of ([\w\s\+\#\.\-]+)',
            r'skilled (?:in|at|with) ([\w\s\+\#\.\-]+)',
            r'familiar with ([\w\s\+\#\.\-]+)',
            r'(\d+)\+?\s*years? of (?:experience )?(?:in|with) ([\w\s\+\#\.\-]+)'
        ]
    
    def get_all_skills(self) -> List[str]:
        """Get flattened list of all skills"""
        all_skills = []
        for skills in self.skills.values():
            all_skills.extend(skills)
        return all_skills
    
    def get_category_for_skill(self, skill: str) -> Optional[str]:
        """Find which category a skill belongs to"""
        skill_lower = skill.lower()
        for category, skills in self.skills.items():
            if any(s.lower() == skill_lower for s in skills):
                return category
        return 'other'


class TextPreprocessor:
    """Advanced text preprocessing for skill extraction"""
    
    def __init__(self):
        self.nlp = self._load_spacy_model()
        self._customize_stop_words()
    
    def _load_spacy_model(self):
        """Load spaCy model with error handling"""
        try:
            return spacy.load("en_core_web_sm")
        except OSError:
            st.warning("⚠️ Installing spaCy model...")
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
            return spacy.load("en_core_web_sm")
    
    def _customize_stop_words(self):
        """Customize stop words for skill extraction"""
        programming_langs = {'c', 'r', 'go', 'd', 'f'}
        for lang in programming_langs:
            self.nlp.Defaults.stop_words.discard(lang)
    
    def preprocess(self, text: str) -> Dict:
        """Complete preprocessing pipeline"""
        if not text or not text.strip():
            return {'success': False, 'error': 'Empty text'}
        
        try:
            doc = self.nlp(text)
            
            return {
                'success': True,
                'doc': doc,
                'noun_chunks': [chunk.text for chunk in doc.noun_chunks],
                'entities': [(ent.text, ent.label_) for ent in doc.ents],
                'sentences': [sent.text for sent in doc.sents]
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}

# ---------------- Logging ----------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------- Document Extraction ----------------
class TextExtractor:
    """Extract text from PDF, DOCX, TXT"""
    def extract_pdf(self, file: BytesIO) -> str:
        text = ""
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def extract_docx(self, file: BytesIO) -> str:
        text = ""
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def extract_txt(self, file: BytesIO) -> str:
        return file.read().decode("utf-8", errors="ignore")

    def extract(self, uploaded_file):
        fname = uploaded_file.name.lower()
        if fname.endswith(".pdf"):
            return self.extract_pdf(uploaded_file)
        elif fname.endswith(".docx"):
            return self.extract_docx(uploaded_file)
        elif fname.endswith(".txt"):
            return self.extract_txt(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {uploaded_file.name}")

# ---------------- Text Cleaner ----------------
class TextCleaner:
    """Cleans extracted text"""
    def clean(self, text: str) -> str:
        text = re.sub(r"\s+", " ", text)
        return text.strip()

# ---------------- Document Parser ----------------
class DocumentParser:
    """Parse documents"""
    def parse_file(self, uploaded_file) -> Tuple[str, str]:
        extractor = TextExtractor()
        ext = uploaded_file.name.split('.')[-1].lower()
        text = extractor.extract(uploaded_file)
        return ext, text

# ---------------- Document Processor ----------------
class DocumentProcessor:
    def __init__(self):
        self.parser = DocumentParser()
        self.processed_docs: List[Dict] = []

    def process_document(self, uploaded_file, doc_type: str) -> Dict:
        try:
            ext, text = self.parser.parse_file(uploaded_file)
            cleaner = TextCleaner()
            clean_text = cleaner.clean(text)
            return {
                "filename": uploaded_file.name,
                "document_type": doc_type,
                "extension": ext,
                "content": clean_text,
                "success": True,
                "error": None,
            }
        except Exception as e:
            return {
                "filename": uploaded_file.name,
                "document_type": doc_type,
                "extension": None,
                "content": None,
                "success": False,
                "error": str(e),
            }

    def process_files(self, resumes, jds):
        self.processed_docs = []
        for file in resumes:
            self.processed_docs.append(self.process_document(file, "resume"))
        for file in jds:
            self.processed_docs.append(self.process_document(file, "job_description"))
        return self.processed_docs

    def display_processing_results(self):
        st.header("📄 Processing Results")
        if not self.processed_docs:
            st.info("No documents processed.")
            return
        for doc in self.processed_docs:
            with st.expander(f"{doc['filename']} ({doc['document_type']})"):
                if doc["success"]:
                    st.success("✅ Processed Successfully")
                    st.write(doc["content"][:500] + "..." if len(doc["content"]) > 500 else doc["content"])
                else:
                    st.error(f"❌ Failed: {doc['error']}")

# ---------------- ATS / Resume Analysis ----------------
class ATSAnalyzer:
    """Analyze resume vs JD and self-analysis"""
    @staticmethod
    def analyze(processed_docs):
        resumes = [d for d in processed_docs if d["document_type"]=="resume" and d["success"]]
        jds = [d for d in processed_docs if d["document_type"]=="job_description" and d["success"]]

        if resumes and jds:
            st.subheader("📊 ATS Analysis (Resume vs Job Description)")
            resume_text = " ".join([d["content"] for d in resumes])
            jd_text = " ".join([d["content"] for d in jds])
            vectorizer = TfidfVectorizer(stop_words="english")
            tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            ats_score = round(similarity * 100,2)
            st.success(f"✅ ATS Score: {ats_score}%")

            # Bar chart
            fig, ax = plt.subplots()
            ax.bar(["ATS Score"], [ats_score], color="green" if ats_score>=80 else "orange" if ats_score>=60 else "red")
            ax.set_ylim(0,100)
            ax.set_ylabel("Score (%)")
            st.pyplot(fig)

            # Keyword comparison
            resume_words = set(resume_text.lower().split())
            jd_words = set(jd_text.lower().split())
            common_words = resume_words.intersection(jd_words)
            missing_words = jd_words - resume_words
            st.info(f"🔍 Common Keywords: {len(common_words)} | ❌ Missing Keywords: {len(missing_words)}")
            if missing_words:
                st.write(", ".join(list(missing_words)[:20]))

            # Suggestions
            st.subheader("💡 Suggestions")
            suggestions = []
            if ats_score < 60:
                suggestions.append("Increase keyword overlap with JD.")
            if len(resume_text.split()) < 250:
                suggestions.append("Add more content about projects and experience.")
            if "experience" not in resume_words:
                suggestions.append("Include an 'Experience' section.")
            if "skills" not in resume_words:
                suggestions.append("Include a 'Skills' section with JD keywords.")
            for s in suggestions:
                st.write(f"- {s}")
        elif resumes and not jds:
            st.subheader("📑 Resume Self Analysis")
            resume_text = " ".join([d["content"] for d in resumes])
            st.info(f"📝 Word count: {len(resume_text.split())}")
            st.info(f"🔠 Char count: {len(resume_text)}")
            keywords = ["python","java","sql","machine learning","data","cloud","powerbi"]
            found = [kw for kw in keywords if kw in resume_text.lower()]
            st.success(f"✅ Skills Detected: {', '.join(found) if found else 'No major keywords found'}")
            # Sections check
            sections = ["summary","experience","projects","skills","education","certifications"]
            missing_sections = [s for s in sections if s not in resume_text.lower()]
            if missing_sections:
                st.warning(f"⚠️ Missing Sections: {', '.join(missing_sections)}")
        else:
            st.warning("No resumes or job descriptions available.")
        
# ---------------- ATS / Resume Analysis ----------------
class ATSAnalyzer:
    """Analyze resume vs JD and self-analysis"""
    @staticmethod
    def analyze(processed_docs):
        resumes = [d for d in processed_docs if d["document_type"]=="resume" and d["success"]]
        jds = [d for d in processed_docs if d["document_type"]=="job_description" and d["success"]]

        if resumes and jds:
            st.subheader("📊 ATS Analysis (Resume vs Job Description)")
            resume_text = " ".join([d["content"] for d in resumes])
            jd_text = " ".join([d["content"] for d in jds])
            vectorizer = TfidfVectorizer(stop_words="english")
            tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            ats_score = round(similarity * 100,2)
            st.success(f"✅ ATS Score: {ats_score}%")

            # Bar chart
            fig, ax = plt.subplots()
            ax.bar(["ATS Score"], [ats_score], color="green" if ats_score>=80 else "orange" if ats_score>=60 else "red")
            ax.set_ylim(0,100)
            ax.set_ylabel("Score (%)")
            st.pyplot(fig)

            # Keyword comparison
            resume_words = set(resume_text.lower().split())
            jd_words = set(jd_text.lower().split())
            common_words = resume_words.intersection(jd_words)
            missing_words = jd_words - resume_words
            st.info(f"🔍 Common Keywords: {len(common_words)} | ❌ Missing Keywords: {len(missing_words)}")
            if missing_words:
                st.write(", ".join(list(missing_words)[:20]))

            # Suggestions
            st.subheader("💡 Suggestions")
            suggestions = []
            if ats_score < 60:
                suggestions.append("Increase keyword overlap with JD.")
            if len(resume_text.split()) < 250:
                suggestions.append("Add more content about projects and experience.")
            if "experience" not in resume_words:
                suggestions.append("Include an 'Experience' section.")
            if "skills" not in resume_words:
                suggestions.append("Include a 'Skills' section with JD keywords.")
            for s in suggestions:
                st.write(f"- {s}")
        elif resumes and not jds:
            st.subheader("📑 Resume Self Analysis")
            resume_text = " ".join([d["content"] for d in resumes])
            st.info(f"📝 Word count: {len(resume_text.split())}")
            st.info(f"🔠 Char count: {len(resume_text)}")
            keywords = ["python","java","sql","machine learning","data","cloud","powerbi"]
            found = [kw for kw in keywords if kw in resume_text.lower()]
            st.success(f"✅ Skills Detected: {', '.join(found) if found else 'No major keywords found'}")
            # Sections check
            sections = ["summary","experience","projects","skills","education","certifications"]
            missing_sections = [s for s in sections if s not in resume_text.lower()]
            if missing_sections:
                st.warning(f"⚠️ Missing Sections: {', '.join(missing_sections)}")
        else:
            st.warning("No resumes or job descriptions available.")



class SkillExtractor:
    """Main skill extraction engine"""
    
    def __init__(self):
        self.skill_db = SkillDatabase()
        self.preprocessor = TextPreprocessor()
        self.logger = self._setup_logger()
    
    def extract_skills(self, text: str, document_type: str = 'resume') -> Dict:
        """Extract skills using multiple methods"""
        try:
            preprocess_result = self.preprocessor.preprocess(text)
            if not preprocess_result['success']:
                return {'success': False, 'error': preprocess_result['error']}
            
            doc = preprocess_result['doc']
            
            # Multiple extraction methods
            keyword_skills = self._extract_by_keywords(text)
            pos_skills = self._extract_by_pos_patterns(doc)
            context_skills = self._extract_by_context(text)
            ner_skills = self._extract_by_ner(preprocess_result['entities'])
            chunk_skills = self._extract_from_noun_chunks(preprocess_result['noun_chunks'])
            
            # Combine and deduplicate
            all_skills = self._combine_and_deduplicate([
                keyword_skills, pos_skills, context_skills, ner_skills, chunk_skills
            ])
            
            # Normalize (expand abbreviations)
            normalized_skills = self._normalize_skills(all_skills)
            
            # Categorize skills
            categorized_skills = self._categorize_skills(normalized_skills)
            
            # Calculate confidence scores
            skill_confidence = self._calculate_confidence(
                normalized_skills,
                [keyword_skills, pos_skills, context_skills, ner_skills, chunk_skills]
            )
            
            return {
                'success': True,
                'all_skills': normalized_skills,
                'categorized_skills': categorized_skills,
                'skill_confidence': skill_confidence,
                'extraction_methods': {
                    'keyword_matching': len(keyword_skills),
                    'pos_patterns': len(pos_skills),
                    'context_based': len(context_skills),
                    'ner': len(ner_skills),
                    'noun_chunks': len(chunk_skills)
                },
                'statistics': {
                    'total_skills': len(normalized_skills),
                    'technical_skills': sum(len(skills) for cat, skills in categorized_skills.items() if cat != 'soft_skills'),
                    'soft_skills': len(categorized_skills.get('soft_skills', []))
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _extract_by_keywords(self, text: str) -> Set[str]:
        """Extract skills by keyword matching"""
        found_skills = set()
        text_lower = text.lower()
        
        for skill in self.skill_db.get_all_skills():
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill)
        
        return found_skills
    
    def _extract_by_pos_patterns(self, doc) -> Set[str]:
        """Extract skills using POS patterns"""
        found_skills = set()
        tokens = list(doc)
        
        # ADJ + NOUN patterns
        for i in range(len(tokens) - 1):
            if tokens[i].pos_ == 'ADJ' and tokens[i+1].pos_ in ['NOUN', 'PROPN']:
                pattern = f"{tokens[i].text} {tokens[i+1].text}"
                if self._is_valid_skill(pattern):
                    found_skills.add(pattern)
        
        # Proper nouns
        for token in doc:
            if token.pos_ == 'PROPN' and self._is_valid_skill(token.text):
                found_skills.add(token.text)
        
        return found_skills
    
    def _extract_by_context(self, text: str) -> Set[str]:
        """Extract skills based on context patterns"""
        found_skills = set()
        
        for pattern in self.skill_db.skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) >= 1:
                    skill_text = match.group(len(match.groups()))
                    skills = self._clean_and_split_skills(skill_text)
                    for skill in skills:
                        if self._is_valid_skill(skill):
                            found_skills.add(skill)
        
        return found_skills
    
    def _extract_by_ner(self, entities: List[Tuple[str, str]]) -> Set[str]:
        """Extract skills from named entities"""
        found_skills = set()
        relevant_labels = ['ORG', 'PRODUCT', 'GPE']
        
        for entity_text, label in entities:
            if label in relevant_labels and self._is_valid_skill(entity_text):
                found_skills.add(entity_text)
        
        return found_skills
    
    def _extract_from_noun_chunks(self, noun_chunks: List[str]) -> Set[str]:
        """Extract skills from noun chunks"""
        found_skills = set()
        
        for chunk in noun_chunks:
            chunk_clean = chunk.strip()
            if self._is_valid_skill(chunk_clean):
                found_skills.add(chunk_clean)
        
        return found_skills
    
    def _is_valid_skill(self, text: str) -> bool:
        """Validate if text is a valid skill"""
        if not text or len(text.strip()) < 2:
            return False
        
        text_clean = text.strip()
        all_skills_lower = [s.lower() for s in self.skill_db.get_all_skills()]
        
        if text_clean.lower() in all_skills_lower:
            return True
        
        for skill in self.skill_db.get_all_skills():
            if skill.lower() in text_clean.lower() or text_clean.lower() in skill.lower():
                if abs(len(skill) - len(text_clean)) <= 3:
                    return True
        
        return False
    
    def _clean_and_split_skills(self, text: str) -> List[str]:
        """Clean and split comma-separated skills"""
        skills = re.split(r'[,;|/&]|\band\b', text)
        cleaned_skills = []
        
        for skill in skills:
            skill_clean = skill.strip()
            skill_clean = re.sub(r'\b(etc|and more)\b', '', skill_clean, flags=re.IGNORECASE).strip()
            if skill_clean and len(skill_clean) > 1:
                cleaned_skills.append(skill_clean)
        
        return cleaned_skills
    
    def _combine_and_deduplicate(self, skill_sets: List[Set[str]]) -> List[str]:
        """Combine and remove duplicates"""
        combined = set()
        for skill_set in skill_sets:
            combined.update(skill_set)
        
        unique_skills = {}
        for skill in combined:
            skill_lower = skill.lower()
            if skill_lower not in unique_skills:
                unique_skills[skill_lower] = skill
        
        return sorted(unique_skills.values())
    
    def _normalize_skills(self, skills: List[str]) -> List[str]:
        """Normalize skill names (expand abbreviations)"""
        normalized = []
        
        for skill in skills:
            if skill.upper() in self.skill_db.abbreviations:
                normalized.append(self.skill_db.abbreviations[skill.upper()])
            else:
                normalized.append(skill)
        
        return sorted(set(normalized))
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills"""
        categorized = defaultdict(list)
        
        for skill in skills:
            category = self.skill_db.get_category_for_skill(skill)
            categorized[category].append(skill)
        
        for category in categorized:
            categorized[category] = sorted(categorized[category])
        
        return dict(categorized)
    
    def _calculate_confidence(self, skills: List[str], method_results: List[Set[str]]) -> Dict[str, float]:
        """Calculate confidence score for each skill"""
        confidence_scores = {}
        
        for skill in skills:
            detection_count = sum(
                1 for method_set in method_results 
                if skill in method_set or skill.lower() in {s.lower() for s in method_set}
            )
            confidence = detection_count / len(method_results)
            confidence_scores[skill] = round(confidence, 2)
        
        return confidence_scores
    
    def _setup_logger(self):
        """Setup logging"""
        logger = logging.getLogger('SkillExtractor')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger


class SentenceBERTEmbedder:
    """Generate and manage BERT embeddings for skills"""
    
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        """Initialize Sentence-BERT model"""
        try:
            self.model = SentenceTransformer(model_name)
            self.skill_embeddings = {}
        except Exception as e:
            st.error(f"❌ Failed to load BERT model: {e}")
            st.info("Installing sentence-transformers...")
            import subprocess
            subprocess.run(["pip", "install", "sentence-transformers"])
            self.model = SentenceTransformer(model_name)
    
    def encode_skills(self, skills: List[str]) -> Dict[str, np.ndarray]:
        """Generate embeddings for skills"""
        if not skills:
            return {}
        
        embeddings = self.model.encode(skills, show_progress_bar=True)
        
        skill_embeddings = {}
        for skill, embedding in zip(skills, embeddings):
            skill_embeddings[skill] = embedding
            self.skill_embeddings[skill] = embedding
        
        return skill_embeddings
    
    def compute_similarity(self, skill1: str, skill2: str) -> float:
        """Compute cosine similarity between two skills"""
        if skill1 not in self.skill_embeddings:
            emb1 = self.model.encode([skill1])[0]
            self.skill_embeddings[skill1] = emb1
        else:
            emb1 = self.skill_embeddings[skill1]
        
        if skill2 not in self.skill_embeddings:
            emb2 = self.model.encode([skill2])[0]
            self.skill_embeddings[skill2] = emb2
        else:
            emb2 = self.skill_embeddings[skill2]
        
        similarity = cosine_similarity([emb1], [emb2])[0][0]
        return float(similarity)
    
    def compute_similarity_matrix(self, skills1: List[str], skills2: List[str]) -> np.ndarray:
        """Compute similarity matrix between two skill sets"""
        embeddings1 = self.model.encode(skills1)
        embeddings2 = self.model.encode(skills2)
        return cosine_similarity(embeddings1, embeddings2)
    
    def find_similar_skills(self, target_skill: str, skill_list: List[str], 
                           threshold: float = 0.7, top_k: int = 5) -> List[Tuple[str, float]]:
        """Find skills similar to target skill"""
        similarities = []
        
        for skill in skill_list:
            if skill.lower() != target_skill.lower():
                sim = self.compute_similarity(target_skill, skill)
                if sim >= threshold:
                    similarities.append((skill, sim))
        
        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:top_k]


class CustomNERTrainer:
    """Train custom NER model for skill detection"""
    
    def __init__(self):
        self.nlp = None
        self.ner = None
    
    def prepare_training_data(self, annotations: List[Dict]) -> List[Tuple]:
        """Convert annotations to spaCy format"""
        training_data = []
        
        for annotation in annotations:
            text = annotation['text']
            entities = []
            
            for skill in annotation['skills']:
                entities.append((skill['start'], skill['end'], skill['label']))
            
            training_data.append((text, {"entities": entities}))
        
        return training_data
    
    def create_blank_model(self):
        """Create blank spaCy model"""
        self.nlp = spacy.blank("en")
        
        if "ner" not in self.nlp.pipe_names:
            self.ner = self.nlp.add_pipe("ner")
        else:
            self.ner = self.nlp.get_pipe("ner")
        
        self.ner.add_label("SKILL")
    
    def train(self, training_data: List[Tuple], n_iterations: int = 30) -> Dict:
        """Train the NER model"""
        if not self.nlp:
            self.create_blank_model()
        
        other_pipes = [pipe for pipe in self.nlp.pipe_names if pipe != "ner"]
        training_stats = {'losses': [], 'iterations': n_iterations}
        
        with self.nlp.disable_pipes(*other_pipes):
            optimizer = self.nlp.begin_training()
            
            for iteration in range(n_iterations):
                random.shuffle(training_data)
                losses = {}
                
                for text, annotations in training_data:
                    doc = self.nlp.make_doc(text)
                    example = Example.from_dict(doc, annotations)
                    self.nlp.update([example], drop=0.5, losses=losses)
                
                training_stats['losses'].append(losses.get('ner', 0))
        
        return training_stats
    
    def predict(self, text: str) -> List[Tuple[str, int, int]]:
        """Predict skills in text"""
        if not self.nlp:
            raise ValueError("Model not trained")
        
        doc = self.nlp(text)
        return [(ent.text, ent.start_char, ent.end_char) for ent in doc.ents if ent.label_ == "SKILL"]


class AnnotationInterface:
    """Interface for creating NER training data"""
    
    def __init__(self):
        if 'training_annotations' not in st.session_state:
            st.session_state.training_annotations = []
        if 'current_skills' not in st.session_state:
            st.session_state.current_skills = []
    
    def create_annotation_ui(self):
        """Create annotation UI"""
        st.subheader("🏷️ Create NER Training Data")
        
        st.markdown("""
        **Instructions:**
        1. Enter text containing skills
        2. Mark skill positions (start/end character indices)
        3. Add to training dataset
        4. Export for model training
        """)
        
        input_text = st.text_area(
            "Enter text to annotate:",
            height=150,
            placeholder="Example: I am a Python developer with 5 years of Machine Learning experience."
        )
        
        if input_text:
            st.markdown("---")
            st.text(input_text)
            
            with st.form("skill_annotation_form"):
                st.markdown("**Add Skill Annotation:**")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    skill_text = st.text_input("Skill text")
                with col2:
                    start_pos = st.number_input("Start position", min_value=0, value=0)
                with col3:
                    end_pos = st.number_input("End position", min_value=0, value=0)
                
                if skill_text and start_pos < end_pos:
                    extracted = input_text[start_pos:end_pos]
                    if extracted.strip():
                        st.info(f"Preview: '{extracted}'")
                
                submitted = st.form_submit_button("➕ Add Skill")
                
                if submitted and skill_text and start_pos < end_pos:
                    st.session_state.current_skills.append({
                        'text': skill_text,
                        'start': start_pos,
                        'end': end_pos,
                        'label': 'SKILL'
                    })
                    st.success(f"✅ Added: {skill_text}")
                    st.rerun()
            
            if st.session_state.current_skills:
                st.markdown("**Skills in current text:**")
                for i, skill in enumerate(st.session_state.current_skills):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.write(f"{i+1}. **{skill['text']}** ({skill['start']}-{skill['end']})")
                    with col2:
                        if st.button("🗑️", key=f"remove_{i}"):
                            st.session_state.current_skills.pop(i)
                            st.rerun()
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Save Annotation", type="primary"):
                    if st.session_state.current_skills:
                        annotation = {
                            'text': input_text,
                            'skills': st.session_state.current_skills.copy(),
                            'timestamp': datetime.now().isoformat()
                        }
                        st.session_state.training_annotations.append(annotation)
                        st.session_state.current_skills = []
                        st.success(f"✅ Saved! Total: {len(st.session_state.training_annotations)}")
                        st.rerun()
            
            with col2:
                if st.button("🔄 Clear Current"):
                    st.session_state.current_skills = []
                    st.rerun()
        
        if st.session_state.training_annotations:
            st.markdown("---")
            st.subheader(f"📚 Training Dataset ({len(st.session_state.training_annotations)} annotations)")
            
            for i, annotation in enumerate(st.session_state.training_annotations):
                with st.expander(f"Annotation {i+1}: {len(annotation['skills'])} skills"):
                    st.text(annotation['text'])
                    st.write("**Skills:**")
                    for skill in annotation['skills']:
                        st.write(f"- {skill['text']} ({skill['start']}-{skill['end']})")
            
            col1, col2 = st.columns(2)
            
            with col1:
                training_json = json.dumps(st.session_state.training_annotations, indent=2)
                st.download_button(
                    "📥 Download Training Data (JSON)",
                    training_json,
                    f"training_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json"
                )
            
            with col2:
                trainer = CustomNERTrainer()
                spacy_format = trainer.prepare_training_data(st.session_state.training_annotations)
                spacy_json = json.dumps(spacy_format, indent=2)
                
                st.download_button(
                    "📥 Download spaCy Format",
                    spacy_json,
                    f"spacy_training_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json"
                )


class SkillVisualizer:
    """Visualize extracted skills"""
    
    @staticmethod
    def create_category_distribution_chart(categorized_skills: Dict[str, List[str]]) -> go.Figure:
        """Create pie chart for category distribution"""
        category_names = {
            'programming_languages': 'Programming Languages',
            'web_frameworks': 'Web Frameworks',
            'databases': 'Databases',
            'ml_ai': 'ML/AI',
            'ml_frameworks': 'ML Frameworks',
            'cloud_platforms': 'Cloud Platforms',
            'devops_tools': 'DevOps Tools',
            'version_control': 'Version Control',
            'testing': 'Testing',
            'soft_skills': 'Soft Skills',
            'other': 'Other'
        }
        
        categories = []
        counts = []
        
        for category, skills in categorized_skills.items():
            if skills:
                categories.append(category_names.get(category, category.replace('_', ' ').title()))
                counts.append(len(skills))
        
        fig = go.Figure(data=[go.Pie(
            labels=categories,
            values=counts,
            hole=0.3,
            textposition='auto',
            textinfo='label+percent+value'
        )])
        
        fig.update_layout(title="Skill Distribution by Category", height=500)
        return fig
    
    @staticmethod
    def create_top_skills_chart(skills: List[str], confidence_scores: Dict[str, float], top_n: int = 15) -> go.Figure:
        """Create bar chart for top skills"""
        sorted_skills = sorted(confidence_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]
        
        skill_names = [skill for skill, _ in sorted_skills]
        confidences = [score * 100 for _, score in sorted_skills]
        
        fig = go.Figure(data=[go.Bar(
            x=confidences,
            y=skill_names,
            orientation='h',
            marker=dict(
                color=confidences,
                colorscale='Viridis',
                colorbar=dict(title="Confidence %")
            ),
            text=[f"{conf:.0f}%" for conf in confidences],
            textposition='auto'
        )])
        
        fig.update_layout(
            title=f"Top {top_n} Skills by Confidence Score",
            xaxis_title="Confidence Score (%)",
            yaxis_title="Skills",
            height=600,
            yaxis=dict(autorange="reversed")
        )
        
        return fig
    
    @staticmethod
    def create_extraction_methods_chart(extraction_methods: Dict[str, int]) -> go.Figure:
        """Create bar chart for extraction methods"""
        method_names = {
            'keyword_matching': 'Keyword Matching',
            'pos_patterns': 'POS Patterns',
            'context_based': 'Context-Based',
            'ner': 'Named Entity Recognition',
            'noun_chunks': 'Noun Chunks'
        }
        
        methods = [method_names.get(m, m) for m in extraction_methods.keys()]
        counts = list(extraction_methods.values())
        
        fig = go.Figure(data=[go.Bar(
            x=methods,
            y=counts,
            marker_color=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd'],
            text=counts,
            textposition='auto'
        )])
        
        fig.update_layout(
            title="Skills Detected by Each Extraction Method",
            xaxis_title="Extraction Method",
            yaxis_title="Number of Skills Found",
            height=400
        )
        
        return fig


class CompleteSkillExtractionApp:
    """Complete Milestone 2 application"""
    
    def __init__(self):
        self.skill_extractor = SkillExtractor()
        self.visualizer = SkillVisualizer()
        self.bert_embedder = SentenceBERTEmbedder()
        self.ner_trainer = CustomNERTrainer()
        self.annotator = AnnotationInterface()
        
        # Initialize session state
        if 'extraction_results' not in st.session_state:
            st.session_state.extraction_results = None
        if 'skill_embeddings' not in st.session_state:
            st.session_state.skill_embeddings = None
        if 'trained_ner' not in st.session_state:
            st.session_state.trained_ner = None
    
    def run(self):
        """Run the complete application"""
        st.title("🎯 AI Skill Gap Analyzer - Complete Milestone 2")
        st.markdown("### Advanced Skill Extraction with BERT Embeddings & Custom NER")
        
        tabs = st.tabs([
            "📄 Extract Skills",
            "🧠 BERT Embeddings",
            "🏋️ Train Custom NER",
            "🏷️ Annotate Data",
            "📊 Visualizations",
            "📥 Export"
        ])
        
        with tabs[0]:
            self._skill_extraction_tab()
        
        with tabs[1]:
            self._bert_embeddings_tab()
        
        with tabs[2]:
            self._ner_training_tab()
        
        with tabs[3]:
            self._annotation_tab()
        
        with tabs[4]:
            self._visualization_tab()
        
        with tabs[5]:
            self._export_tab()
    
    def _skill_extraction_tab(self):
        """Skill extraction interface"""
        st.header("Extract Skills from Text")
        
        input_method = st.radio(
            "Choose input method:",
            ["Paste Text", "Upload File"],
            horizontal=True
        )
        
        text_input = ""
        doc_type = "resume"
        
        if input_method == "Paste Text":
            col1, col2 = st.columns([3, 1])
            
            with col1:
                text_input = st.text_area(
                    "Paste resume or job description text:",
                    height=300,
                    placeholder="Paste your resume or job description here..."
                )
            
            with col2:
                doc_type = st.selectbox("Document Type:", ["resume", "job_description"])
        
        else:
            uploaded_file = st.file_uploader("Upload document", type=['txt'])
            doc_type = st.selectbox("Document Type:", ["resume", "job_description"])
            
            if uploaded_file:
                text_input = uploaded_file.getvalue().decode("utf-8", errors="replace")
        
        if st.button("🔍 Extract Skills", type="primary", use_container_width=True):
            if text_input:
                with st.spinner("Extracting skills..."):
                    result = self.skill_extractor.extract_skills(text_input, doc_type)
                    
                    if result['success']:
                        st.session_state.extraction_results = result
                        st.success(f"✅ Successfully extracted {result['statistics']['total_skills']} skills!")
                        self._display_extraction_results(result)
                    else:
                        st.error(f"❌ Extraction failed: {result.get('error', 'Unknown error')}")
            else:
                st.warning("⚠️ Please provide text input")
    
    def _display_extraction_results(self, result: Dict):
        """Display extraction results"""
        st.subheader("📊 Extraction Results")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Skills", result['statistics']['total_skills'])
        with col2:
            st.metric("Technical Skills", result['statistics']['technical_skills'])
        with col3:
            st.metric("Soft Skills", result['statistics']['soft_skills'])
        with col4:
            avg_confidence = sum(result['skill_confidence'].values()) / len(result['skill_confidence']) if result['skill_confidence'] else 0
            st.metric("Avg Confidence", f"{avg_confidence:.0%}")
        
        st.subheader("🏷️ Categorized Skills")
        
        categorized = result['categorized_skills']
        category_items = list(categorized.items())
        
        for i in range(0, len(category_items), 2):
            cols = st.columns(2)
            
            for j, col in enumerate(cols):
                if i + j < len(category_items):
                    category, skills = category_items[i + j]
                    
                    with col:
                        category_display = category.replace('_', ' ').title()
                        st.markdown(f"**{category_display}** ({len(skills)})")
                        
                        skill_html = ""
                        for skill in skills[:10]:
                            confidence = result['skill_confidence'].get(skill, 0)
                            color = "tech-skill" if category != "soft_skills" else "soft-skill"
                            skill_html += f'<span class="skill-tag {color}" title="Confidence: {confidence:.0%}">{skill}</span>'
                        
                        if len(skills) > 10:
                            skill_html += f'<span class="skill-tag">+{len(skills) - 10} more</span>'
                        
                        st.markdown(skill_html, unsafe_allow_html=True)
                        st.markdown("")
        
        with st.expander("🔧 Extraction Methods Used"):
            methods_df = pd.DataFrame([
                {'Method': method.replace('_', ' ').title(), 'Skills Found': count}
                for method, count in result['extraction_methods'].items()
            ])
            st.dataframe(methods_df, use_container_width=True)
    
    def _bert_embeddings_tab(self):
        """BERT embeddings interface"""
        st.header("🧠 Skill Embeddings with Sentence-BERT")
        
        st.markdown("""
        **Sentence-BERT** creates semantic embeddings for skills, enabling:
        - Similarity computation between skills
        - Semantic skill matching
        - Finding related skills
        """)
        
        if not st.session_state.extraction_results:
            st.info("👆 Extract skills first to generate embeddings")
            return
        
        result = st.session_state.extraction_results
        skills = result['all_skills']
        
        if st.button("🚀 Generate BERT Embeddings", type="primary"):
            with st.spinner("Generating embeddings..."):
                embeddings = self.bert_embedder.encode_skills(skills)
                st.session_state.skill_embeddings = embeddings
                st.success(f"✅ Generated embeddings for {len(skills)} skills!")
        
        if st.session_state.skill_embeddings:
            st.subheader("🔍 Skill Similarity Calculator")
            
            col1, col2 = st.columns(2)
            
            with col1:
                skill1 = st.selectbox("Select first skill:", skills, key="sim_skill1")
            
            with col2:
                skill2 = st.selectbox("Select second skill:", skills, key="sim_skill2")
            
            if st.button("Calculate Similarity"):
                similarity = self.bert_embedder.compute_similarity(skill1, skill2)
                
                st.metric(
                    "Similarity Score",
                    f"{similarity:.2%}",
                    delta=f"{'High' if similarity > 0.7 else 'Medium' if similarity > 0.4 else 'Low'} similarity"
                )
                
                fig = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=similarity * 100,
                    domain={'x': [0, 1], 'y': [0, 1]},
                    title={'text': "Similarity"},
                    gauge={
                        'axis': {'range': [0, 100]},
                        'bar': {'color': "darkblue"},
                        'steps': [
                            {'range': [0, 40], 'color': "lightgray"},
                            {'range': [40, 70], 'color': "gray"},
                            {'range': [70, 100], 'color': "lightgreen"}
                        ]
                    }
                ))
                
                st.plotly_chart(fig, use_container_width=True)
            
            st.subheader("🎯 Find Similar Skills")
            
            target_skill = st.selectbox("Select target skill:", skills, key="target_skill")
            threshold = st.slider("Similarity threshold:", 0.0, 1.0, 0.7, 0.05)
            
            if st.button("Find Similar Skills"):
                similar_skills = self.bert_embedder.find_similar_skills(
                    target_skill,
                    [s for s in skills if s != target_skill],
                    threshold=threshold,
                    top_k=10
                )
                
                if similar_skills:
                    st.success(f"Found {len(similar_skills)} similar skills:")
                    
                    df = pd.DataFrame(similar_skills, columns=['Skill', 'Similarity'])
                    df['Similarity'] = df['Similarity'].apply(lambda x: f"{x:.2%}")
                    st.dataframe(df, use_container_width=True)
                    
                    fig = go.Figure(data=[go.Bar(
                        x=[s[1] for s in similar_skills],
                        y=[s[0] for s in similar_skills],
                        orientation='h',
                        marker_color='lightblue'
                    )])
                    
                    fig.update_layout(
                        title=f"Skills Similar to '{target_skill}'",
                        xaxis_title="Similarity Score",
                        yaxis_title="Skill",
                        height=400
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning(f"No skills found with similarity >= {threshold:.0%}")
            
            st.subheader("📊 Skill Similarity Matrix")
            
            if st.button("Generate Similarity Matrix"):
                with st.spinner("Computing similarities..."):
                    similarity_matrix = self.bert_embedder.compute_similarity_matrix(
                        skills[:20],
                        skills[:20]
                    )
                    
                    fig = go.Figure(data=go.Heatmap(
                        z=similarity_matrix,
                        x=skills[:20],
                        y=skills[:20],
                        colorscale='Viridis',
                        text=similarity_matrix,
                        texttemplate='%{text:.2f}',
                        textfont={"size": 8}
                    ))
                    
                    fig.update_layout(
                        title="Skill Similarity Heatmap (Top 20 Skills)",
                        height=700,
                        width=800
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
    
    def _ner_training_tab(self):
        """Custom NER training interface"""
        st.header("🏋️ Train Custom NER Model")
        
        st.markdown("""
        Train a custom spaCy NER model to detect skills in text.
        
        **Steps:**
        1. Load training data (use Annotate Data tab)
        2. Configure training parameters
        3. Train the model
        4. Test the model
        """)
        
        st.subheader("1️⃣ Load Training Data")
        
        training_source = st.radio(
            "Training data source:",
            ["Use Annotated Data", "Upload JSON File"],
            horizontal=True
        )
        
        training_data = None
        
        if training_source == "Use Annotated Data":
            if st.session_state.get('training_annotations'):
                st.success(f"✅ {len(st.session_state.training_annotations)} annotations available")
                training_data = self.ner_trainer.prepare_training_data(
                    st.session_state.training_annotations
                )
            else:
                st.warning("⚠️ No annotations found. Use 'Annotate Data' tab first.")
        else:
            uploaded_file = st.file_uploader("Upload training data (JSON)", type=['json'])
            if uploaded_file:
                try:
                    annotations = json.load(uploaded_file)
                    training_data = self.ner_trainer.prepare_training_data(annotations)
                    st.success(f"✅ Loaded {len(training_data)} training examples")
                except Exception as e:
                    st.error(f"❌ Error loading file: {e}")
        
        if training_data:
            st.subheader("2️⃣ Configure Training")
            
            col1, col2 = st.columns(2)
            
            with col1:
                n_iterations = st.number_input(
                    "Number of iterations:",
                    min_value=10,
                    max_value=100,
                    value=30,
                    step=10
                )
            
            with col2:
                st.info(f"Training examples: {len(training_data)}")
            
            st.subheader("3️⃣ Train Model")
            
            if st.button("🚀 Start Training", type="primary"):
                with st.spinner("Training model..."):
                    progress_bar = st.progress(0)
                    
                    try:
                        self.ner_trainer.create_blank_model()
                        training_stats = self.ner_trainer.train(training_data, n_iterations=n_iterations)
                        
                        st.session_state.trained_ner = self.ner_trainer
                        st.session_state.training_stats = training_stats
                        
                        progress_bar.progress(100)
                        st.success("✅ Training complete!")
                        
                        fig = go.Figure()
                        fig.add_trace(go.Scatter(
                            x=list(range(1, len(training_stats['losses']) + 1)),
                            y=training_stats['losses'],
                            mode='lines+markers',
                            name='Training Loss'
                        ))
                        
                        fig.update_layout(
                            title="Training Loss Over Iterations",
                            xaxis_title="Iteration",
                            yaxis_title="Loss",
                            height=400
                        )
                        
                        st.plotly_chart(fig, use_container_width=True)
                        
                    except Exception as e:
                        st.error(f"❌ Training failed: {e}")
            
            if st.session_state.get('trained_ner'):
                st.subheader("4️⃣ Test Model")
                
                test_text = st.text_area(
                    "Enter text to test model:",
                    placeholder="Example: I am proficient in Python, Java, and Machine Learning."
                )
                
                if st.button("🧪 Test"):
                    if test_text:
                        try:
                            predictions = st.session_state.trained_ner.predict(test_text)
                            
                            if predictions:
                                st.success(f"✅ Found {len(predictions)} skills:")
                                for skill, start, end in predictions:
                                    st.markdown(f"- **{skill}** (position {start}-{end})")
                            else:
                                st.warning("No skills detected")
                        except Exception as e:
                            st.error(f"❌ Prediction failed: {e}")
    
    def _annotation_tab(self):
        """Annotation interface"""
        self.annotator.create_annotation_ui()
    
    def _visualization_tab(self):
        """Visualization interface"""
        if not st.session_state.extraction_results:
            st.info("👆 Please extract skills first in the 'Extract Skills' tab")
            return
        
        result = st.session_state.extraction_results
        
        st.header("📊 Skill Analysis Visualizations")
        
        st.subheader("Skill Distribution by Category")
        fig_pie = self.visualizer.create_category_distribution_chart(result['categorized_skills'])
        st.plotly_chart(fig_pie, use_container_width=True)
        
        st.subheader("Top Skills by Confidence Score")
        top_n = st.slider("Number of top skills to display:", 5, 30, 15)
        fig_bar = self.visualizer.create_top_skills_chart(
            result['all_skills'],
            result['skill_confidence'],
            top_n
        )
        st.plotly_chart(fig_bar, use_container_width=True)
        
        st.subheader("Extraction Methods Comparison")
        fig_methods = self.visualizer.create_extraction_methods_chart(result['extraction_methods'])
        st.plotly_chart(fig_methods, use_container_width=True)
        
        st.subheader("📋 Detailed Skill Table")
        
        detailed_data = []
        for skill in result['all_skills']:
            category = self.skill_extractor.skill_db.get_category_for_skill(skill)
            confidence = result['skill_confidence'].get(skill, 0)
            
            detailed_data.append({
                'Skill': skill,
                'Category': category.replace('_', ' ').title(),
                'Confidence': f"{confidence:.0%}",
                'Confidence Score': confidence
            })
        
        df = pd.DataFrame(detailed_data)
        
        col1, col2 = st.columns(2)
        with col1:
            categories = ['All'] + sorted(df['Category'].unique().tolist())
            selected_category = st.selectbox("Filter by category:", categories)
        
        with col2:
            min_confidence = st.slider("Minimum confidence:", 0.0, 1.0, 0.0, 0.1)
        
        filtered_df = df.copy()
        if selected_category != 'All':
            filtered_df = filtered_df[filtered_df['Category'] == selected_category]
        filtered_df = filtered_df[filtered_df['Confidence Score'] >= min_confidence]
        
        st.dataframe(
            # filtered_df[['Skill', 'Category', 'Confidence']].sort_values('Confidence Score', ascending=False),
            filtered_df[['Skill', 'Category', 'Confidence']].sort_values('Confidence', ascending=False),

            use_container_width=True,
            hide_index=True
        )
        
        st.caption(f"Showing {len(filtered_df)} of {len(df)} skills")
    
    def _export_tab(self):
        """Export interface"""
        if not st.session_state.extraction_results:
            st.info("👆 Please extract skills first in the 'Extract Skills' tab")
            return
        
        result = st.session_state.extraction_results
        
        st.header("📥 Export Extracted Skills")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            csv_data = self._create_csv_export(result)
            st.download_button(
                label="📄 Download CSV",
                data=csv_data,
                file_name=f"extracted_skills_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with col2:
            json_data = self._create_json_export(result)
            st.download_button(
                label="📋 Download JSON",
                data=json_data,
                file_name=f"extracted_skills_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
        
        with col3:
            report_data = self._create_text_report(result)
            st.download_button(
                label="📑 Download Report",
                data=report_data,
                file_name=f"skill_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
    
    def _create_csv_export(self, result: Dict) -> str:
        """Create CSV export"""
        data = []
        
        for skill in result['all_skills']:
            category = self.skill_extractor.skill_db.get_category_for_skill(skill)
            confidence = result['skill_confidence'].get(skill, 0)
            
            data.append({
                'Skill': skill,
                'Category': category,
                'Confidence': confidence,
                'Type': 'Soft Skill' if category == 'soft_skills' else 'Technical Skill'
            })
        
        df = pd.DataFrame(data)
        return df.to_csv(index=False)
    
    def _create_json_export(self, result: Dict) -> str:
        """Create JSON export"""
        export_data = {
            'extraction_timestamp': datetime.now().isoformat(),
            'statistics': result['statistics'],
            'skills': {
                'all_skills': result['all_skills'],
                'categorized_skills': result['categorized_skills'],
                'skill_confidence': result['skill_confidence']
            },
            'extraction_methods': result['extraction_methods']
        }
        
        return json.dumps(export_data, indent=2)
    
    def _create_text_report(self, result: Dict) -> str:
        """Create formatted text report"""
        report = []
        report.append("=" * 80)
        report.append("SKILL EXTRACTION REPORT")
        report.append("=" * 80)
        report.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"\nTotal Skills Extracted: {result['statistics']['total_skills']}")
        report.append(f"Technical Skills: {result['statistics']['technical_skills']}")
        report.append(f"Soft Skills: {result['statistics']['soft_skills']}")
        report.append("\n" + "-" * 80)
        report.append("\nCATEGORIZED SKILLS")
        report.append("-" * 80)
        
        for category, skills in sorted(result['categorized_skills'].items()):
            if skills:
                report.append(f"\n{category.replace('_', ' ').title()} ({len(skills)}):")
                for skill in skills:
                    confidence = result['skill_confidence'].get(skill, 0)
                    report.append(f"  • {skill} (Confidence: {confidence:.0%})")
        
        report.append("\n" + "=" * 80)
        report.append("END OF REPORT")
        report.append("=" * 80)
        
        return "\n".join(report)



def main():
    """Main application entry point"""
    try:
        app = CompleteSkillExtractionApp()
        app.run()
        
        st.title("📄 Complete Skill Extraction & ATS Analyzer")

        # File upload section
        st.sidebar.header("Upload Documents")
        resumes = st.sidebar.file_uploader(
            "Upload Resumes (PDF/DOCX/TXT)", type=["pdf","docx","txt"], accept_multiple_files=True
        )
        jds = st.sidebar.file_uploader(
            "Upload Job Descriptions (PDF/DOCX/TXT)", type=["pdf","docx","txt"], accept_multiple_files=True
        )

        # Process uploaded files
        processed_docs = []
        if resumes or jds:
            processor = DocumentProcessor()
            processed_docs = processor.process_files(resumes, jds)
            processor.display_processing_results()

            # ATS Analysis
            ATSAnalyzer.analyze(processed_docs)

        # Sidebar info
        with st.sidebar:
            st.header("ℹ️ Milestone 2 Complete")
            st.markdown("""
            **✅ All Features Implemented:**
            
            1. **Skill Extraction (NLP)**
               - spaCy pipeline
               - Multi-method extraction
               - 5 extraction techniques
            
            2. **BERT Embeddings**
               - Sentence-BERT
               - Similarity computation
               - Semantic matching
            
            3. **Custom NER Training**
               - Model training
               - Testing interface
               - Training visualization
            
            4. **Annotation Interface**
               - Training data creation
               - Export functionality
            
            5. **Visualizations**
               - Interactive charts
               - Distribution analysis
               - Similarity heatmaps
            
            6. **Export Options**
               - CSV, JSON, Text reports
            """)
            
            if processed_docs:
                total_skills = sum(len(d.get('content','').split()) for d in processed_docs if d["success"])
                st.header("📊 Current Stats")
                st.metric("Documents Processed", len(processed_docs))
                st.metric("Total Words Processed", total_skills)

    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.exception(e)


if __name__ == "__main__":
    main()







































































































































































































































































































































































































































































# """
# Complete AI Skill Gap Analyzer
# Milestone 1 & 2 Integrated Pipeline

# Requirements:
# pip install streamlit PyPDF2 python-docx pandas scikit-learn matplotlib plotly spacy sentence-transformers
# Run: streamlit run complete_skill_gap_analyzer.py
# """

# import streamlit as st
# import PyPDF2
# import docx
# import pandas as pd
# import re
# import logging
# import random
# import json
# from typing import Dict, List, Tuple
# from io import BytesIO
# from datetime import datetime

# # ML / Analysis Imports
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# import matplotlib.pyplot as plt
# import plotly.graph_objects as go
# import spacy
# from spacy.training import Example
# from sentence_transformers import SentenceTransformer
# from collections import defaultdict

# # ---------------- Logging ----------------
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # ---------------- Document Extraction ----------------
# class TextExtractor:
#     """Extract text from PDF, DOCX, TXT"""
#     def extract_pdf(self, file: BytesIO) -> str:
#         text = ""
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text() or ""
#         return text

#     def extract_docx(self, file: BytesIO) -> str:
#         text = ""
#         doc = docx.Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#         return text

#     def extract_txt(self, file: BytesIO) -> str:
#         return file.read().decode("utf-8", errors="ignore")

#     def extract(self, uploaded_file):
#         fname = uploaded_file.name.lower()
#         if fname.endswith(".pdf"):
#             return self.extract_pdf(uploaded_file)
#         elif fname.endswith(".docx"):
#             return self.extract_docx(uploaded_file)
#         elif fname.endswith(".txt"):
#             return self.extract_txt(uploaded_file)
#         else:
#             raise ValueError(f"Unsupported file type: {uploaded_file.name}")

# # ---------------- Text Cleaner ----------------
# class TextCleaner:
#     """Cleans extracted text"""
#     def clean(self, text: str) -> str:
#         text = re.sub(r"\s+", " ", text)
#         return text.strip()

# # ---------------- Document Parser ----------------
# class DocumentParser:
#     """Parse documents"""
#     def parse_file(self, uploaded_file) -> Tuple[str, str]:
#         extractor = TextExtractor()
#         ext = uploaded_file.name.split('.')[-1].lower()
#         text = extractor.extract(uploaded_file)
#         return ext, text

# # ---------------- Document Processor ----------------
# class DocumentProcessor:
#     def __init__(self):
#         self.parser = DocumentParser()
#         self.processed_docs: List[Dict] = []

#     def process_document(self, uploaded_file, doc_type: str) -> Dict:
#         try:
#             ext, text = self.parser.parse_file(uploaded_file)
#             cleaner = TextCleaner()
#             clean_text = cleaner.clean(text)
#             return {
#                 "filename": uploaded_file.name,
#                 "document_type": doc_type,
#                 "extension": ext,
#                 "content": clean_text,
#                 "success": True,
#                 "error": None,
#             }
#         except Exception as e:
#             return {
#                 "filename": uploaded_file.name,
#                 "document_type": doc_type,
#                 "extension": None,
#                 "content": None,
#                 "success": False,
#                 "error": str(e),
#             }

#     def process_files(self, resumes, jds):
#         self.processed_docs = []
#         for file in resumes:
#             self.processed_docs.append(self.process_document(file, "resume"))
#         for file in jds:
#             self.processed_docs.append(self.process_document(file, "job_description"))
#         return self.processed_docs

#     def display_processing_results(self):
#         st.header("📄 Processing Results")
#         if not self.processed_docs:
#             st.info("No documents processed.")
#             return
#         for doc in self.processed_docs:
#             with st.expander(f"{doc['filename']} ({doc['document_type']})"):
#                 if doc["success"]:
#                     st.success("✅ Processed Successfully")
#                     st.write(doc["content"][:500] + "..." if len(doc["content"]) > 500 else doc["content"])
#                 else:
#                     st.error(f"❌ Failed: {doc['error']}")

# # ---------------- ATS / Resume Analysis ----------------
# class ATSAnalyzer:
#     """Analyze resume vs JD and self-analysis"""
#     @staticmethod
#     def analyze(processed_docs):
#         resumes = [d for d in processed_docs if d["document_type"]=="resume" and d["success"]]
#         jds = [d for d in processed_docs if d["document_type"]=="job_description" and d["success"]]

#         if resumes and jds:
#             st.subheader("📊 ATS Analysis (Resume vs Job Description)")
#             resume_text = " ".join([d["content"] for d in resumes])
#             jd_text = " ".join([d["content"] for d in jds])
#             vectorizer = TfidfVectorizer(stop_words="english")
#             tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
#             similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
#             ats_score = round(similarity * 100,2)
#             st.success(f"✅ ATS Score: {ats_score}%")

#             # Bar chart
#             fig, ax = plt.subplots()
#             ax.bar(["ATS Score"], [ats_score], color="green" if ats_score>=80 else "orange" if ats_score>=60 else "red")
#             ax.set_ylim(0,100)
#             ax.set_ylabel("Score (%)")
#             st.pyplot(fig)

#             # Keyword comparison
#             resume_words = set(resume_text.lower().split())
#             jd_words = set(jd_text.lower().split())
#             common_words = resume_words.intersection(jd_words)
#             missing_words = jd_words - resume_words
#             st.info(f"🔍 Common Keywords: {len(common_words)} | ❌ Missing Keywords: {len(missing_words)}")
#             if missing_words:
#                 st.write(", ".join(list(missing_words)[:20]))

#             # Suggestions
#             st.subheader("💡 Suggestions")
#             suggestions = []
#             if ats_score < 60:
#                 suggestions.append("Increase keyword overlap with JD.")
#             if len(resume_text.split()) < 250:
#                 suggestions.append("Add more content about projects and experience.")
#             if "experience" not in resume_words:
#                 suggestions.append("Include an 'Experience' section.")
#             if "skills" not in resume_words:
#                 suggestions.append("Include a 'Skills' section with JD keywords.")
#             for s in suggestions:
#                 st.write(f"- {s}")
#         elif resumes and not jds:
#             st.subheader("📑 Resume Self Analysis")
#             resume_text = " ".join([d["content"] for d in resumes])
#             st.info(f"📝 Word count: {len(resume_text.split())}")
#             st.info(f"🔠 Char count: {len(resume_text)}")
#             keywords = ["python","java","sql","machine learning","data","cloud","powerbi"]
#             found = [kw for kw in keywords if kw in resume_text.lower()]
#             st.success(f"✅ Skills Detected: {', '.join(found) if found else 'No major keywords found'}")
#             # Sections check
#             sections = ["summary","experience","projects","skills","education","certifications"]
#             missing_sections = [s for s in sections if s not in resume_text.lower()]
#             if missing_sections:
#                 st.warning(f"⚠️ Missing Sections: {', '.join(missing_sections)}")
#         else:
#             st.warning("No resumes or job descriptions available.")

# # ---------------- Skill Extractor ----------------
# class SkillExtractor:
#     """Simple keyword-based skill extraction (for demonstration)"""
#     def __init__(self):
#         self.skill_db = {
#             "python":"programming_languages", "java":"programming_languages",
#             "sql":"databases","machine learning":"ml_ai","data":"ml_ai",
#             "cloud":"cloud_platforms","powerbi":"ml_ai"
#         }
#     def extract_skills(self, text, doc_type="resume"):
#         skills_found = []
#         confidence = {}
#         for skill in self.skill_db.keys():
#             if skill.lower() in text.lower():
#                 skills_found.append(skill)
#                 confidence[skill] = random.uniform(0.6,1.0)
#         categorized = defaultdict(list)
#         for skill in skills_found:
#             categorized[self.skill_db[skill]].append(skill)
#         return {
#             "success": True,
#             "all_skills": skills_found,
#             "categorized_skills": categorized,
#             "skill_confidence": confidence,
#             "statistics":{
#                 "total_skills": len(skills_found),
#                 "technical_skills": len(skills_found),
#                 "soft_skills": 0
#             },
#             "extraction_methods":{"keyword_matching":len(skills_found)}
#         }

# # ---------------- Sentence BERT Embeddings ----------------
# class SentenceBERTEmbedder:
#     def __init__(self):
#         self.model = SentenceTransformer('all-MiniLM-L6-v2')
#         self.embeddings = {}

#     def encode_skills(self, skills: List[str]):
#         self.embeddings = {skill:self.model.encode(skill) for skill in skills}
#         return self.embeddings

#     def compute_similarity(self, skill1, skill2):
#         from numpy import dot
#         from numpy.linalg import norm
#         v1, v2 = self.embeddings[skill1], self.embeddings[skill2]
#         return dot(v1,v2)/(norm(v1)*norm(v2))

# # ---------------- Custom NER Trainer ----------------
# class CustomNERTrainer:
#     def __init__(self):
#         self.nlp = None
#         self.ner = None

#     def create_blank_model(self):
#         self.nlp = spacy.blank("en")
#         self.ner = self.nlp.add_pipe("ner")
#         self.ner.add_label("SKILL")

#     def prepare_training_data(self, annotations):
#         spacy_data = []
#         for ann in annotations:
#             entities = [(s['start'],s['end'],s['label']) for s in ann['skills']]
#             spacy_data.append((ann['text'], {"entities":entities}))
#         return spacy_data

#     def train(self, training_data: List[Tuple], n_iterations:int=30):
#         if not self.nlp:
#             self.create_blank_model()
#         other_pipes = [pipe for pipe in self.nlp.pipe_names if pipe!="ner"]
#         stats = {'losses':[], 'iterations': n_iterations}
#         with self.nlp.disable_pipes(*other_pipes):
#             optimizer = self.nlp.begin_training()
#             for i in range(n_iterations):
#                 random.shuffle(training_data)
#                 losses = {}
#                 for text, ann in training_data:
#                     doc = self.nlp.make_doc(text)
#                     example = Example.from_dict(doc, ann)
#                     self.nlp.update([example], drop=0.5, losses=losses)
#                 stats['losses'].append(losses.get('ner',0))
#         return stats

#     def predict(self, text: str):
#         if not self.nlp:
#             raise ValueError("Model not trained")
#         doc = self.nlp(text)
#         return [(ent.text, ent.start_char, ent.end_char) for ent in doc.ents if ent.label_=="SKILL"]

# # ---------------- Annotation Interface ----------------
# class AnnotationInterface:
#     def __init__(self):
#         if 'training_annotations' not in st.session_state:
#             st.session_state.training_annotations = []
#         if 'current_skills' not in st.session_state:
#             st.session_state.current_skills = []

#     def create_annotation_ui(self):
#         st.subheader("🏷️ Annotate Skills for NER Training")
#         text = st.text_area("Enter text to annotate:")
#         if text:
#             with st.form("annotation_form"):
#                 skill_text = st.text_input("Skill text")
#                 start = st.number_input("Start position", 0)
#                 end = st.number_input("End position",0)
#                 submitted = st.form_submit_button("➕ Add Skill")
#                 if submitted and skill_text and start<end:
#                     st.session_state.current_skills.append({'text':skill_text,'start':start,'end':end,'label':'SKILL'})
#                     st.success(f"Added: {skill_text}")
#                     st.rerun()
#             if st.session_state.current_skills:
#                 st.write(st.session_state.current_skills)
#                 if st.button("✅ Save Annotation"):
#                     st.session_state.training_annotations.append({'text':text,'skills':st.session_state.current_skills})
#                     st.session_state.current_skills=[]

# # ---------------- Visualization ----------------
# class Visualizer:
#     def create_category_distribution_chart(self, categorized_skills):
#         fig = go.Figure()
#         for cat, skills in categorized_skills.items():
#             fig.add_trace(go.Bar(name=cat, x=skills, y=[1]*len(skills)))
#         fig.update_layout(barmode='stack', title="Skill Distribution by Category", height=400)
#         return fig

#     def create_top_skills_chart(self, skills, confidence, top_n=10):
#         sorted_skills = sorted(skills, key=lambda x: confidence.get(x,0), reverse=True)[:top_n]
#         fig = go.Figure([go.Bar(x=sorted_skills, y=[confidence[s] for s in sorted_skills])])
#         fig.update_layout(title=f"Top {top_n} Skills by Confidence", height=400)
#         return fig

# # ---------------- Main Application ----------------
# class CompleteSkillExtractionApp:
#     def __init__(self):
#         self.doc_processor = DocumentProcessor()
#         self.skill_extractor = SkillExtractor()
#         self.bert_embedder = SentenceBERTEmbedder()
#         self.ner_trainer = CustomNERTrainer()
#         self.annotator = AnnotationInterface()
#         self.visualizer = Visualizer()

#     def run(self):
#         st.set_page_config(page_title="AI Skill Gap Analyzer", layout="wide")
#         st.title("💼 AI Skill Gap Analyzer")

#         tabs = st.tabs(["1️⃣ Upload & Process","2️⃣ ATS / Resume Analysis","3️⃣ Extract Skills",
#                         "4️⃣ Custom NER Trainer","5️⃣ Annotation","6️⃣ Visualizations","7️⃣ Export"])

#         with tabs[0]:
#             self.upload_tab()
#         with tabs[1]:
#             if st.session_state.get('processed_docs'):
#                 ATSAnalyzer.analyze(st.session_state['processed_docs'])
#             else:
#                 st.info("Upload documents first.")
#         with tabs[2]:
#             if st.session_state.get('processed_docs'):
#                 self.skills_tab()
#             else:
#                 st.info("Upload documents first.")
#         with tabs[3]:
#             self.ner_tab()
#         with tabs[4]:
#             self.annotator.create_annotation_ui()
#         with tabs[5]:
#             self.visualization_tab()
#         with tabs[6]:
#             self.export_tab()

#     def upload_tab(self):
#         with st.sidebar:
#             st.header("Upload Documents")
#         resumes = st.file_uploader("Upload Resumes", type=["pdf","docx","txt"], accept_multiple_files=True)
#         jds = st.file_uploader("Upload Job Descriptions", type=["pdf","docx","txt"], accept_multiple_files=True)
#         if st.button("Process Documents"):
#             processed = self.doc_processor.process_files(resumes, jds)
#             st.session_state['processed_docs'] = processed
#             st.success("Documents processed successfully!")
#         if st.session_state.get('processed_docs'):
#             self.doc_processor.display_processing_results()

#     def skills_tab(self):
#         st.header("📌 Extract Skills")
#         extracted_results = []
#         for doc in st.session_state['processed_docs']:
#             if doc['success']:
#                 res = self.skill_extractor.extract_skills(doc['content'], doc['document_type'])
#                 extracted_results.append(res)
#         st.session_state['extraction_results'] = extracted_results
#         st.success("Skills extracted successfully!")
#         st.write(extracted_results)

#     def ner_tab(self):
#         st.header("🧪 Custom NER Trainer")
#         n_iterations = st.number_input("Training Iterations", 5,100,20)
#         if st.button("Train NER Model"):
#             try:
#                 self.ner_trainer.create_blank_model()
#                 # Dummy training using annotations
#                 training_data = st.session_state.get('training_annotations',[])
#                 spacy_data = self.ner_trainer.prepare_training_data(training_data)
#                 training_stats = self.ner_trainer.train(spacy_data, n_iterations=n_iterations)
#                 st.session_state['trained_ner'] = self.ner_trainer
#                 st.session_state['training_stats'] = training_stats
#                 st.success("NER Training Completed")
#                 fig = go.Figure([go.Scatter(x=list(range(1,len(training_stats['losses'])+1)),
#                                             y=training_stats['losses'],mode='lines+markers',name='Loss')])
#                 fig.update_layout(title="Training Loss",xaxis_title="Iteration",yaxis_title="Loss",height=400)
#                 st.plotly_chart(fig)
#             except Exception as e:
#                 st.error(f"Training failed: {e}")
#         if st.session_state.get('trained_ner'):
#             test_text = st.text_area("Enter text to test NER model:")
#             if st.button("Test NER"):
#                 try:
#                     predictions = st.session_state.trained_ner.predict(test_text)
#                     if predictions:
#                         st.success(f"Found {len(predictions)} skills")
#                         for skill,start,end in predictions:
#                             st.write(f"- {skill} ({start}-{end})")
#                     else:
#                         st.warning("No skills detected")
#                 except Exception as e:
#                     st.error(f"Prediction failed: {e}")

#     def visualization_tab(self):
#         if not st.session_state.get('extraction_results'):
#             st.info("Extract skills first.")
#             return
#         st.header("📊 Skill Visualizations")
#         for res in st.session_state['extraction_results']:
#             fig = self.visualizer.create_category_distribution_chart(res['categorized_skills'])
#             st.plotly_chart(fig,use_container_width=True)
#             top_fig = self.visualizer.create_top_skills_chart(res['all_skills'],res['skill_confidence'])
#             st.plotly_chart(top_fig,use_container_width=True)

#     def export_tab(self):
#         if not st.session_state.get('extraction_results'):
#             st.info("Extract skills first.")
#             return
#         st.header("📥 Export Skills")
#         result = st.session_state['extraction_results'][0]
#         df = pd.DataFrame([{'Skill':k,'Confidence':v} for k,v in result['skill_confidence'].items()])
#         csv_data = df.to_csv(index=False)
#         json_data = json.dumps(result,indent=2)
#         st.download_button("Download CSV", data=csv_data, file_name="skills.csv", mime="text/csv")
#         st.download_button("Download JSON", data=json_data, file_name="skills.json", mime="application/json")

# # ---------------- Run App ----------------
# def main():
#     try:
#         app = CompleteSkillExtractionApp()
#         app.run()
#     except Exception as e:
#         st.error(f"Application error: {str(e)}")
#         st.exception(e)

# if __name__=="__main__":
#     main()
