import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from pathlib import Path
import PyPDF2
import docx
import re
from io import BytesIO
from typing import Dict, List, Optional
from datetime import datetime, timedelta

import logging
import json
from dataclasses import dataclass
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import spacy



# # ==================== CUSTOM STYLING ====================
# st.markdown("""
#     <style>
#         /* Background Gradient */
#         .stApp {
#             background: linear-gradient(to right, #e8f5e9, #e0f7fa);
#         }

#         /* Title Styling */
#         h1, h2, h3 {
#             color: #004d40 !important;
#             text-shadow: 1px 1px 2px #b2dfdb;
#         }

#         /* Buttons */
#         .stButton button {
#             background-color: #26a69a;
#             color: white;
#             border-radius: 12px;
#             border: none;
#             padding: 8px 20px;
#             font-size: 16px;
#             transition: 0.3s;
#         }
#         .stButton button:hover {
#             background-color: #00796b;
#             transform: scale(1.05);
#         }

#         /* Upload Box */
#         .upload-box {
#             border: 2px dashed #80cbc4;
#             padding: 25px;
#             border-radius: 15px;
#             background-color: #f9fbe7;
#         }

#         /* Skill Tag */
#         .skill-tag {
#             display: inline-block;
#             background-color: #81c784;
#             color: white;
#             padding: 6px 12px;
#             margin: 4px;
#             border-radius: 15px;
#             font-size: 14px;
#             box-shadow: 1px 1px 4px rgba(0,0,0,0.2);
#         }

#         /* Footer */
#         footer {
#             visibility: hidden;
#         }
#     </style>
# """, unsafe_allow_html=True)

# # ==================== HEADER ====================
# st.markdown("""
# <div style='text-align: center;'>
#     <img src="https://cdn-icons-png.flaticon.com/512/2721/2721298.png" width="100">
#     <h1 style='color: #00695c;'>AI Skill Gap Analyzer</h1>
#     <p style='font-size:18px;'>Empowering Careers through AI-driven Skill Insights 💼🤖</p>
# </div>
# """, unsafe_allow_html=True)





# ==================== CUSTOM STYLING ====================
st.markdown("""
    <style>
        /* Background Gradient Animation */
        .stApp {
            background: linear-gradient(120deg, #e8f5e9, #e0f7fa, #f1f8e9);
            background-size: 300% 300%;
            animation: gradientShift 10s ease infinite alternate;
        }

        @keyframes gradientShift {
            0% { background-position: left; }
            100% { background-position: right; }
        }

        /* Header & Title Animations */
        h1, h2, h3 {
            color: #004d40 !important;
            text-shadow: 1px 1px 2px #b2dfdb;
            animation: fadeInUp 1.5s ease-in-out;
        }

        p {
            animation: fadeIn 2s ease-in-out;
        }

        @keyframes fadeInUp {
            0% { opacity: 0; transform: translateY(30px); }
            100% { opacity: 1; transform: translateY(0); }
        }

        @keyframes fadeIn {
            from { opacity: 0; }
            to { opacity: 1; }
        }

        /* Floating Logo Animation */
        .floating-img {
            animation: float 3s ease-in-out infinite;
        }
        @keyframes float {
            0% { transform: translateY(0px); }
            50% { transform: translateY(-10px); }
            100% { transform: translateY(0px); }
        }

        /* Buttons */
        .stButton button {
            background-color: #26a69a;
            color: white;
            border-radius: 12px;
            border: none;
            padding: 8px 20px;
            font-size: 16px;
            transition: 0.3s;
            box-shadow: 0 0 0px #26a69a;
        }
        .stButton button:hover {
            background-color: #00796b;
            transform: scale(1.05);
            box-shadow: 0 0 15px #4db6ac;
        }

        /* Upload Box */
        .upload-box {
            border: 2px dashed #80cbc4;
            padding: 25px;
            border-radius: 15px;
            background-color: #f9fbe7;
            animation: fadeIn 2s ease-in-out;
        }

        /* Skill Tags */
        .skill-tag {
            display: inline-block;
            background-color: #81c784;
            color: white;
            padding: 6px 12px;
            margin: 4px;
            border-radius: 15px;
            font-size: 14px;
            box-shadow: 1px 1px 4px rgba(0,0,0,0.2);
            transition: transform 0.3s ease, background-color 0.3s ease;
        }
        .skill-tag:hover {
            transform: scale(1.1);
            background-color: #4caf50;
        }

        /* ==================== TAB ANIMATION ==================== */
        div[data-baseweb="tab-list"] {
            display: flex;
            justify-content: center;
            gap: 10px;
            margin-top: 20px;
            flex-wrap: wrap;
            animation: fadeIn 1.2s ease-in-out;
        }

        div[data-baseweb="tab"] {
            background-color: rgba(255, 255, 255, 0.85);
            color: #004d40 !important;
            font-weight: 600;
            border-radius: 10px;
            padding: 10px 20px;
            transition: all 0.4s ease-in-out;
            box-shadow: 0px 2px 5px rgba(0,0,0,0.15);
        }

        div[data-baseweb="tab"]:hover {
            background-color: #26a69a !important;
            color: white !important;
            transform: translateY(-4px);
            box-shadow: 0px 6px 12px rgba(0,0,0,0.2);
        }

        div[data-baseweb="tab"][aria-selected="true"] {
            background-color: #00796b !important;
            color: white !important;
            animation: tab-glow 1s ease-in-out infinite alternate;
        }

        @keyframes tab-glow {
            0% { box-shadow: 0 0 6px #4db6ac; }
            100% { box-shadow: 0 0 20px #004d40; }
        }

        /* Fade-in effect for tab content */
        section.main > div {
            animation: fadeContent 0.8s ease-in-out;
        }

        @keyframes fadeContent {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }

        /* Hide footer */
        footer { visibility: hidden; }
    </style>
""", unsafe_allow_html=True)

# ==================== HEADER ====================
st.markdown("""
<div style='text-align: center; animation: fadeIn 1.5s ease-in-out;'>
    <img src="https://cdn-icons-png.flaticon.com/512/2721/2721298.png" width="100" class="floating-img">
    <h1 style='color: #00695c;'>AI Skill Gap Analyzer</h1>
    <p style='font-size:18px;'>Empowering Careers through AI-driven Skill Insights 💼🤖</p>
</div>
""", unsafe_allow_html=True)







# --- Embedded lightweight skill database & extractor (replaces milestone2.SkillExtractor) ---
class SkillDatabase:
    def __init__(self):
        # smaller curated list for offline use — extend as needed
        self.skills = {
            'programming_languages': ['Python', 'Java', 'JavaScript', 'C++', 'C#', 'R'],
            'web_frameworks': ['Django', 'Flask', 'React', 'Angular', 'Node.js'],
            'databases': ['MySQL', 'PostgreSQL', 'MongoDB', 'Redis'],
            'ml_ai': ['Machine Learning', 'Deep Learning', 'NLP', 'Computer Vision'],
            'ml_frameworks': ['TensorFlow', 'PyTorch', 'Scikit-learn'],
            'cloud_platforms': ['AWS', 'Azure', 'GCP'],
            'devops_tools': ['Docker', 'Kubernetes'],
            'soft_skills': ['Communication', 'Leadership', 'Teamwork']
        }

    def get_all_skills(self):
        all_skills = []
        for v in self.skills.values():
            all_skills.extend(v)
        return all_skills

    def get_category_for_skill(self, skill: str):
        s = skill.lower()
        for cat, skills in self.skills.items():
            if any(sk.lower() == s for sk in skills):
                return cat
        return 'other'


class SkillExtractor:
    """Lightweight skill extractor using keyword matching"""
    def __init__(self):
        self.skill_db = SkillDatabase()
        self.logger = logging.getLogger('SkillExtractor')

    def extract_skills(self, text: str, document_type: str = 'resume') -> dict:
        if not text:
            return {'success': False, 'error': 'No text provided', 'all_skills': [], 'categorized_skills': {}, 'skill_confidence': {}}

        text_lower = text.lower()
        found = set()
        for skill in self.skill_db.get_all_skills():
            if skill.lower() in text_lower:
                found.add(skill)

        normalized = sorted(found)
        categorized = {}
        for s in normalized:
            cat = self.skill_db.get_category_for_skill(s)
            categorized.setdefault(cat, []).append(s)

        # simple confidence: 1.0 for keyword match
        confidences = {s: 1.0 for s in normalized}

        stats = {
            'total_skills': len(normalized),
            'technical_skills': sum(len(v) for k, v in categorized.items() if k != 'soft_skills'),
            'soft_skills': len(categorized.get('soft_skills', []))
        }

        return {
            'success': True,
            'all_skills': normalized,
            'categorized_skills': categorized,
            'skill_confidence': confidences,
            'extraction_methods': {},
            'statistics': stats
        }


# --- Lightweight embedding & gap analysis (replaces milestone3 components) ---
@dataclass
class SkillMatch:
    jd_skill: str
    resume_skill: str
    similarity: float
    category: str
    confidence_level: str
    priority: str = 'MEDIUM'


@dataclass
class GapAnalysisResult:
    matched_skills: list
    partial_matches: list
    missing_skills: list
    overall_score: float
    category_scores: dict
    similarity_matrix: any
    resume_skills: list
    jd_skills: list

    def get_statistics(self):
        total = len(self.jd_skills)
        return {
            'total_required_skills': total,
            'matched_count': len(self.matched_skills),
            'partial_count': len(self.partial_matches),
            'missing_count': len(self.missing_skills),
            'match_percentage': (len(self.matched_skills) / total * 100) if total > 0 else 0,
            'overall_score': self.overall_score * 100
        }


class SentenceBERTEncoder:
    """Encoder that uses SentenceTransformer when available, falls back to TF-IDF vectors"""
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        self.logger = logging.getLogger('Encoder')
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer(model_name)
            self.use_transformer = True
            self.logger.info('Loaded SentenceTransformer model')
        except Exception:
            self.model = TfidfVectorizer().fit
            self.use_transformer = False
            self.logger.info('SentenceTransformer not available — will use TF-IDF fallback')

    def encode_skills(self, skills: list, show_progress: bool = False):
        if not skills:
            return np.zeros((0, 1))
        if self.use_transformer:
            embeddings = self.model.encode(skills)
            return np.array(embeddings)
        else:
            # TF-IDF fallback — vectorize skill tokens
            vect = TfidfVectorizer()
            X = vect.fit_transform(skills).toarray()
            return np.array(X)


class SimilarityCalculator:
    def compute_similarity_matrix(self, resume_embeddings, jd_embeddings):
        if resume_embeddings.size == 0 or jd_embeddings.size == 0:
            return np.zeros((len(resume_embeddings), len(jd_embeddings)))
        return cosine_similarity(resume_embeddings, jd_embeddings)


class SkillGapAnalyzer:
    def __init__(self, encoder: SentenceBERTEncoder = None, calculator: SimilarityCalculator = None,
                 strong_threshold: float = 0.8, partial_threshold: float = 0.5):
        self.encoder = encoder or SentenceBERTEncoder()
        self.calculator = calculator or SimilarityCalculator()
        self.strong_threshold = strong_threshold
        self.partial_threshold = partial_threshold
        self.logger = logging.getLogger('GapAnalyzer')

    def analyze(self, resume_skills: list, jd_skills: list, skill_categories: dict = None):
        if not resume_skills or not jd_skills:
            return GapAnalysisResult([], [], jd_skills or [], 0.0, {}, np.zeros((0,0)), resume_skills or [], jd_skills or [])

        re_emb = self.encoder.encode_skills(resume_skills)
        jd_emb = self.encoder.encode_skills(jd_skills)

        sim_matrix = self.calculator.compute_similarity_matrix(re_emb, jd_emb)

        matched = []
        partial = []
        missing = []
        
        # Calculate skill importance based on frequency in job description
        skill_importance = {}
        if skill_categories:
            for skill in jd_skills:
                # Check if skill is a core technical skill
                cat = skill_categories.get(skill, '').lower()
                is_core_tech = any(tech in cat for tech in ['programming', 'database', 'framework', 'ml_', 'cloud'])
                # High priority for core technical skills mentioned multiple times
                mentions = sum(1 for s in jd_skills if s.lower() == skill.lower())
                skill_importance[skill] = 'HIGH' if (mentions > 1 or is_core_tech) else 'MEDIUM'

        for jdx, jd in enumerate(jd_skills):
            if sim_matrix.size == 0:
                best_sim = 0.0
                best_idx = -1
            else:
                best_idx = int(np.argmax(sim_matrix[:, jdx]))
                best_sim = float(sim_matrix[best_idx, jdx])

            resume_skill = resume_skills[best_idx] if best_idx >= 0 and resume_skills else ''
            cat = (skill_categories.get(jd) if skill_categories else None) or 'other'
            
            # Determine priority based on similarity and importance
            importance = skill_importance.get(jd, 'MEDIUM')
            
            if best_sim >= self.strong_threshold:
                matched.append(SkillMatch(jd, resume_skill, best_sim, cat, 'HIGH'))
            elif best_sim >= self.partial_threshold:
                priority = 'HIGH' if importance == 'HIGH' else 'MEDIUM'
                partial.append(SkillMatch(jd, resume_skill, best_sim, cat, 'MEDIUM', priority=priority))
            else:
                # Missing skills get high priority if they are important or technical
                priority = importance
                missing.append(SkillMatch(jd, '', 0.0, cat, 'LOW', priority=priority))

        overall_score = float(np.mean(sim_matrix.max(axis=0))) if sim_matrix.size else 0.0
        category_scores = {}
        return GapAnalysisResult(matched, partial, missing, overall_score, category_scores, sim_matrix, resume_skills, jd_skills)


class GapVisualizer:
    @staticmethod
    def create_similarity_heatmap(similarity_matrix, resume_skills, jd_skills):
        max_display = 20
        display_resume = resume_skills[:max_display]
        display_jd = jd_skills[:max_display]
        dm = similarity_matrix[:max_display, :max_display] if similarity_matrix.size else np.zeros((len(display_resume), len(display_jd)))
        fig = go.Figure(data=go.Heatmap(z=dm, x=display_jd, y=display_resume, colorscale='RdYlGn', zmid=0.5))
        fig.update_layout(title='Skill Similarity Heatmap')
        return fig

    @staticmethod
    def create_match_distribution_pie(analysis_result: GapAnalysisResult):
        stats = analysis_result.get_statistics()
        fig = go.Figure(data=go.Pie(labels=['Strong', 'Partial', 'Missing'], values=[stats['matched_count'], stats['partial_count'], stats['missing_count']], hole=0.4))
        fig.update_layout(title='Match Distribution')
        return fig

    @staticmethod
    def create_gap_priority_chart(missing_skills):
        if not missing_skills:
            return go.Figure()
        skills = [s.jd_skill for s in missing_skills]
        priorities = [1 if s.priority == 'HIGH' else 0.5 if s.priority == 'MEDIUM' else 0.2 for s in missing_skills]
        fig = go.Figure(data=go.Bar(x=skills, y=priorities))
        fig.update_layout(title='Gap Priority')
        return fig

    @staticmethod
    def create_overall_score_gauge(overall_score: float):
        fig = go.Figure(go.Indicator(mode='gauge+number', value=overall_score * 100, title={'text': 'Overall Match Score'}, gauge={'axis': {'range': [0, 100]}}))
        return fig


class ReportGenerator:
    def generate_text_report(self, analysis_result: GapAnalysisResult) -> str:
        stats = analysis_result.get_statistics()
        lines = []
        lines.append('SKILL GAP REPORT')
        lines.append(f"Overall match: {stats['overall_score']:.1f}%")
        lines.append(f"Matched: {stats['matched_count']}")
        lines.append(f"Partial: {stats['partial_count']}")
        lines.append(f"Missing: {stats['missing_count']}")
        for m in analysis_result.matched_skills:
            lines.append(f"MATCH: {m.jd_skill} -> {m.resume_skill} ({m.similarity:.2f})")
        return '\n'.join(lines)

    def generate_csv_report(self, analysis_result: GapAnalysisResult) -> str:
        rows = []
        for m in analysis_result.matched_skills + analysis_result.partial_matches + analysis_result.missing_skills:
            rows.append({'jd_skill': m.jd_skill, 'resume_skill': m.resume_skill, 'similarity': m.similarity, 'priority': m.priority})
        return pd.DataFrame(rows).to_csv(index=False)

    def generate_json_report(self, analysis_result: GapAnalysisResult) -> str:
        data = {
            'statistics': analysis_result.get_statistics(),
            'matched': [m.__dict__ for m in analysis_result.matched_skills],
            'partial': [m.__dict__ for m in analysis_result.partial_matches],
            'missing': [m.__dict__ for m in analysis_result.missing_skills]
        }
        return json.dumps(data, indent=2)


class LearningPathGenerator:
    def __init__(self):
        # Base time estimates for different skill categories
        self.category_estimates = {
            'programming_languages': {'base_time': '8-12', 'complexity': 'HIGH'},
            'web_frameworks': {'base_time': '6-8', 'complexity': 'MEDIUM'},
            'databases': {'base_time': '4-6', 'complexity': 'MEDIUM'},
            'ml_ai': {'base_time': '12-16', 'complexity': 'HIGH'},
            'ml_frameworks': {'base_time': '8-10', 'complexity': 'HIGH'},
            'cloud_platforms': {'base_time': '6-8', 'complexity': 'MEDIUM'},
            'devops_tools': {'base_time': '4-6', 'complexity': 'MEDIUM'},
            'soft_skills': {'base_time': '4-6', 'complexity': 'LOW'}
        }
        
        # Specific skill overrides for common skills
        self.db = {
            'Python': {'time_estimate': '6-10', 'complexity': 'MEDIUM', 
                      'resources': ['Automate the Boring Stuff', 'Official Tutorial']},
            'Java': {'time_estimate': '8-12', 'complexity': 'HIGH',
                    'resources': ['Oracle Java Tutorial', 'Spring Documentation']},
            'Machine Learning': {'time_estimate': '12-16', 'complexity': 'HIGH',
                               'resources': ['Andrew Ng Course', 'Fast.ai']},
            'TensorFlow': {'time_estimate': '8-10', 'complexity': 'HIGH',
                          'resources': ['TensorFlow Documentation', 'Coursera']},
            'Docker': {'time_estimate': '3-5', 'complexity': 'MEDIUM',
                      'resources': ['Docker Docs', 'Docker in Practice']},
            'AWS': {'time_estimate': '8-12', 'complexity': 'HIGH',
                   'resources': ['AWS Training', 'Cloud Practitioner']}
        }

    def _get_time_estimate(self, skill_name, category, priority):
        # Get base estimate from specific skill or category
        if skill_name in self.db:
            base = self.db[skill_name]['time_estimate']
            complexity = self.db[skill_name]['complexity']
        else:
            cat_info = self.category_estimates.get(category, {'base_time': '4-6', 'complexity': 'MEDIUM'})
            base = cat_info['base_time']
            complexity = cat_info['complexity']
        
        # Parse the range
        low, high = map(int, base.split('-'))
        
        # Adjust based on priority and complexity
        if priority == 'HIGH':
            if complexity == 'HIGH':
                return f"{high}-{high + 4}"  # More time for complex high-priority skills
            return f"{low}-{high}"  # Full range for high priority
        elif priority == 'MEDIUM':
            if complexity == 'HIGH':
                return f"{high - 2}-{high}"  # Upper range for complex skills
            return f"{low + 2}-{high - 2}"  # Middle range for medium priority
        else:  # LOW priority
            return f"{low}-{low + 2}"  # Lower range for low priority

    def generate_path(self, missing_skills, current_skills):
        plan = []
        for m in missing_skills:
            skill_name = m.jd_skill
            category = m.category if hasattr(m, 'category') else 'other'
            
            # Get time estimate based on skill, category and priority
            time_estimate = self._get_time_estimate(skill_name, category, m.priority)
            
            # Get resources
            if skill_name in self.db:
                resources = self.db[skill_name]['resources']
            else:
                cat_name = category.replace('_', ' ').title()
                resources = [f'Search {cat_name} tutorials and documentation']
            
            plan.append({
                'skill': skill_name,
                'time_estimate': time_estimate,
                'priority': m.priority,
                'category': category,
                'resources': resources
            })
        return plan


# Configure page
st.set_page_config(
    page_title="AI Skill Gap Analyzer",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Apply light-green background and black text styling (user request)
st.markdown("""
<style>
  /* Page background and primary text color */
  html, body, .stApp, .block-container, .css-18e3th9 { background-color: #e8f5e9 !important; color: #000 !important; }
  .sidebar .sidebar-content { background-color: #e8f5e9 !important; color: #000 !important; }
  h1, h2, h3, h4, h5, h6, p, span, div { color: #000 !important; }

  /* Skill tag styles (user-provided) */
  .skill-tag {
    display: inline-block;
    padding: 4px 8px;
    margin: 4px;
    border-radius: 6px;
    background: #ffffff;
    color: #000;
    box-shadow: 0 1px 2px rgba(0,0,0,0.06);
    font-weight: 600;
  }

  .tech-skill {
    background: #c8e6c9; /* light green */
    color: #000;
    border: 1px solid #a5d6a7;
  }

  .soft-skill {
    background: #fff9c4; /* light yellow */
    color: #000;
    border: 1px solid #fff59d;
  }

  /* Change button colors in Score & Report and Learning Path tabs */
  div[data-testid="stHorizontalBlock"] button[kind="secondary"],
  [data-testid="baseButton-secondary"] {
    background-color: #ff4444 !important;
    color: white !important;
    border-color: #cc0000 !important;
  }

  /* Ensure download buttons in last tabs are red */
  div[data-baseweb="button"] button,
  .stDownloadButton button {
    background-color: #ff4444 !important;
    color: white !important;
    border-color: #cc0000 !important;
  }
</style>
""", unsafe_allow_html=True)

class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        
    def process_document(self, file, use_ocr: bool = False, remove_pii: bool = False) -> str:
        """Process uploaded document and extract text"""
        if file is None:
            return ""
            
        file_ext = Path(file.name).suffix[1:].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
            
        content = file.read()
        
        if file_ext == 'pdf':
            return self._process_pdf(content, use_ocr)
        elif file_ext == 'docx':
            return self._process_docx(content)
        else:
            return content.decode('utf-8')
            
    def _process_pdf(self, content: bytes, use_ocr: bool) -> str:
        """Extract text from PDF"""
        pdf = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
        return text.strip()
        
    def _process_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(BytesIO(content))
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return "\n".join(text).strip()

class SkillGapApp:
    def __init__(self):
        # Initialize components
        self.doc_processor = DocumentProcessor()
        self.skill_extractor = SkillExtractor()
        self.analyzer = SkillGapAnalyzer(
            encoder=SentenceBERTEncoder(),
            calculator=SimilarityCalculator()
        )
        self.visualizer = GapVisualizer()
        self.report_gen = ReportGenerator()
        self.learning_path_gen = LearningPathGenerator()
        
        # Initialize session state
        if 'resume_text' not in st.session_state:
            st.session_state.resume_text = None
        if 'jd_text' not in st.session_state:
            st.session_state.jd_text = None
        if 'extracted_skills' not in st.session_state:
            st.session_state.extracted_skills = {}
        if 'analysis_results' not in st.session_state:
            st.session_state.analysis_results = None

    def run(self):
        # st.title("🎯 AI Skill Gap Analyzer")
        
        # Main navigation
        tabs = st.tabs([
            "📄 Upload Documents",
            "🔍 Extract Skills",
            "📊 Analysis",
            "📈 Visualization",
            "🎯 Score & Report",
            "📚 Learning Path"
        ])
        
        # Sidebar settings
        with st.sidebar:
            st.header("⚙️ Settings")
            similarity_threshold = st.slider(
                "Similarity Threshold",
                0.0, 1.0, 0.6,
                help="Minimum similarity score for skill matching"
            )
            confidence_threshold = st.slider(
                "Confidence Threshold",
                0.0, 1.0, 0.6,
                help="Minimum confidence score for skill extraction"
            )
            
            # Processing options
            st.subheader("Document Processing")
            use_ocr = st.checkbox("Use OCR for PDFs", True)
            remove_pii = st.checkbox("Remove personal info", True)

        # Upload Documents Tab
        with tabs[0]:
            self.upload_tab(use_ocr, remove_pii)

        # Extract Skills Tab
        with tabs[1]:
            self.extract_tab(confidence_threshold)

        # Analysis Tab
        with tabs[2]:
            self.analysis_tab(similarity_threshold)

        # Visualization Tab
        with tabs[3]:
            self.visualization_tab()

        # Score & Report Tab
        with tabs[4]:
            self.score_tab()

        # Learning Path Tab
        with tabs[5]:
            self.learning_tab()

    def upload_tab(self, use_ocr: bool, remove_pii: bool):
        st.header("📄 Document Upload")
        st.write("Upload your resume and job description to begin the analysis")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Resume Upload")
            resume_file = st.file_uploader(
                "Upload Resume",
                type=['pdf', 'docx', 'txt'],
                key="resume_upload"
            )
            if resume_file:
                try:
                    resume_text = self.doc_processor.process_document(
                        resume_file,
                        use_ocr=use_ocr,
                        remove_pii=remove_pii
                    )
                    st.session_state.resume_text = resume_text
                    st.success("✅ Resume processed successfully!")
                    with st.expander("View Processed Text"):
                        st.text_area("Resume Content", resume_text, height=200)
                except Exception as e:
                    st.error(f"Error processing resume: {str(e)}")

        with col2:
            st.subheader("Job Description Upload")
            jd_file = st.file_uploader(
                "Upload Job Description",
                type=['pdf', 'docx', 'txt'],
                key="jd_upload"
            )
            if jd_file:
                try:
                    jd_text = self.doc_processor.process_document(
                        jd_file,
                        use_ocr=use_ocr,
                        remove_pii=remove_pii
                    )
                    st.session_state.jd_text = jd_text
                    st.success("✅ Job Description processed successfully!")
                    with st.expander("View Processed Text"):
                        st.text_area("JD Content", jd_text, height=200)
                except Exception as e:
                    st.error(f"Error processing job description: {str(e)}")

    def extract_tab(self, confidence_threshold: float):
        st.header("🔍 Extract Skills")

        if not (st.session_state.resume_text and st.session_state.jd_text):
            st.warning("⚠️ Please upload both resume and job description first")
            return

        if st.button("Extract Skills", type="primary"):
            with st.spinner("Extracting skills..."):
                # Extract skills from resume
                resume_result = self.skill_extractor.extract_skills(st.session_state.resume_text)
                resume_skills = resume_result['all_skills'] if resume_result['success'] else []
                
                # Extract skills from job description
                jd_result = self.skill_extractor.extract_skills(st.session_state.jd_text)
                jd_skills = jd_result['all_skills'] if jd_result['success'] else []

                # Store results
                st.session_state.extracted_skills = {
                    'resume': resume_skills,
                    'jd': jd_skills
                }

                st.success("✅ Skills extracted successfully!")

        if st.session_state.extracted_skills:
            # Display statistics
            st.subheader("📊 Extraction Statistics")
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("### Resume Skills")
                resume_skills = st.session_state.extracted_skills['resume']
                st.write(f"**Total Skills Found:** {len(resume_skills)}")
                
                for skill in resume_skills:
                    st.markdown(f"- {skill}")

            with col2:
                st.markdown("### Job Requirements")
                jd_skills = st.session_state.extracted_skills['jd']
                st.write(f"**Total Requirements:** {len(jd_skills)}")
                
                for skill in jd_skills:
                    st.markdown(f"- {skill}")

    def analysis_tab(self, similarity_threshold: float):
        st.header("📊 Skill Gap Analysis")

        if not st.session_state.extracted_skills:
            st.warning("⚠️ Please extract skills first")
            return

        if st.button("Analyze Skill Gaps", type="primary"):
            with st.spinner("Analyzing skill gaps..."):
                # Get skills
                resume_skills = st.session_state.extracted_skills['resume']
                jd_skills = st.session_state.extracted_skills['jd']

                # Perform analysis
                analysis_result = self.analyzer.analyze(
                    resume_skills,
                    jd_skills
                )

                st.session_state.analysis_results = analysis_result
                st.success("✅ Analysis complete!")

        if st.session_state.analysis_results:
            result = st.session_state.analysis_results
            stats = result.get_statistics()

            # Overview metrics
            st.subheader("Analysis Overview")
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric("Overall Match", f"{stats['overall_score']:.1f}%")
            with col2:
                st.metric("Strong Matches", stats['matched_count'])
            with col3:
                st.metric("Partial Matches", stats['partial_count'])
            with col4:
                st.metric("Missing Skills", stats['missing_count'])

            # Detailed results sections
            st.subheader("Detailed Analysis")
            
            # Strong matches
            if result.matched_skills:
                with st.expander("✅ Strong Matches", expanded=True):
                    for match in result.matched_skills:
                        st.info(f"{match.jd_skill} ({match.similarity*100:.1f}% match)")

            # Partial matches
            if result.partial_matches:
                with st.expander("⚠️ Partial Matches"):
                    for match in result.partial_matches:
                        st.warning(
                            f"{match.jd_skill} → {match.resume_skill}"
                            f" ({match.similarity*100:.1f}% match)"
                        )

            # Missing skills
            if result.missing_skills:
                with st.expander("❌ Missing Skills"):
                    for skill in result.missing_skills:
                        st.error(f"{skill.jd_skill} (Priority: {skill.priority})")

    def visualization_tab(self):
        st.header("📈 Visualizations")

        if not st.session_state.analysis_results:
            st.warning("⚠️ Please complete analysis first")
            return

        result = st.session_state.analysis_results

        # Create visualizations
        col1, col2 = st.columns(2)

        with col1:
            # Skill match distribution
            st.subheader("Skill Match Distribution")
            stats = result.get_statistics()
            fig_dist = self.visualizer.create_match_distribution_pie(result)
            st.plotly_chart(fig_dist)

        with col2:
            # Similarity heatmap
            st.subheader("Skill Similarity Matrix")
            fig_heatmap = self.visualizer.create_similarity_heatmap(
                result.similarity_matrix,
                result.resume_skills,
                result.jd_skills
            )
            st.plotly_chart(fig_heatmap)

        # Skill gap analysis
        st.subheader("Skill Gap Analysis")
        fig_gaps = self.visualizer.create_gap_priority_chart(result.missing_skills)
        st.plotly_chart(fig_gaps, use_container_width=True)

        # --- Additional resume-focused visualizations ---
        st.markdown("---")
        st.subheader("Resume Skill Visualizations")

        # Get resume text and skill list
        resume_text = st.session_state.get('resume_text') or ''
        all_known_skills = self.skill_extractor.skill_db.get_all_skills()

        # compute occurrence counts for known skills in the resume text
        counts = {}
        rt_lower = resume_text.lower()
        for s in all_known_skills:
            cnt = rt_lower.count(s.lower())
            if cnt > 0:
                counts[s] = cnt

        if not counts:
            st.info('No known skills found in the resume for additional visualizations.')
            return

        df_sk = pd.DataFrame([{'skill': k, 'count': v, 'category': self.skill_extractor.skill_db.get_category_for_skill(k)} for k, v in counts.items()])

        # Top skills bar chart
        st.markdown('### Top Skills (by frequency)')
        df_bar = df_sk.sort_values('count', ascending=False).head(20)
        fig_bar = px.bar(df_bar, x='skill', y='count', color='count', color_continuous_scale='Blues')
        fig_bar.update_layout(xaxis_tickangle=45, height=350)
        st.plotly_chart(fig_bar, use_container_width=True)

        # Category distribution donut chart
        st.markdown('### Skill Categories (Resume)')
        cat_counts = df_sk.groupby('category').size().reset_index(name='num_skills')
        fig_cat = px.pie(cat_counts, names='category', values='num_skills', hole=0.4)
        st.plotly_chart(fig_cat, use_container_width=True)

        # Bubble chart: frequency vs category (category -> numeric index for y)
        st.markdown('### Skill Frequency vs Category')
        cats = list(df_sk['category'].unique())
        cat_index = {c: i for i, c in enumerate(cats)}
        df_sk['cat_idx'] = df_sk['category'].map(cat_index)
        fig_bubble = px.scatter(df_sk, x='count', y='cat_idx', size='count', hover_name='skill', color='category', labels={'cat_idx': 'Category (index)'}, height=300)
        fig_bubble.update_yaxes(tickmode='array', tickvals=list(cat_index.values()), ticktext=list(cat_index.keys()))
        st.plotly_chart(fig_bubble, use_container_width=True)

        # Small tag cloud (HTML spans sized by count)
        st.markdown('### Tag Cloud')
        max_count = df_sk['count'].max()
        tags_html = []
        for _, row in df_sk.sort_values('count', ascending=False).iterrows():
            size = 12 + int((row['count'] / max_count) * 36)
            tags_html.append(f"<span style='font-size:{size}px; margin:6px; padding:4px; background:#f1f8e9; color:#000; border-radius:6px;'>{row['skill']}</span>")
        st.markdown("" + "".join(tags_html), unsafe_allow_html=True)

            # (Learning timeline moved to its own tab)

    def score_tab(self):
        st.header("🎯 Score & Report")

        if not st.session_state.analysis_results:
            st.warning("⚠️ Please complete analysis first")
            return

        result = st.session_state.analysis_results
        stats = result.get_statistics()

        # Overall score gauge
        fig_gauge = self.visualizer.create_overall_score_gauge(result.overall_score)
        st.plotly_chart(fig_gauge)

        # Export options
        st.subheader("📥 Export Report")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Text report
            text_report = self.report_gen.generate_text_report(result)
            st.download_button(
                "Download Text Report",
                text_report,
                "skill_gap_analysis.txt",
                "text/plain"
            )
        
        with col2:
            # CSV report
            csv_data = self.report_gen.generate_csv_report(result)
            st.download_button(
                "Download CSV Report",
                csv_data,
                "skill_gap_analysis.csv",
                "text/csv"
            )
        
        with col3:
            # JSON report
            json_data = self.report_gen.generate_json_report(result)
            st.download_button(
                "Download JSON Report",
                json_data,
                "skill_gap_analysis.json",
                "application/json"
            )

    def learning_tab(self):
        st.header("📚 Learning Path")

        if not st.session_state.get('analysis_results'):
            st.warning("Please run the Analysis tab to generate missing skills first.")
            return

        result = st.session_state.analysis_results
        if not result.missing_skills:
            st.info("No missing skills detected — nothing to build a learning path from.")
            return

        st.subheader("Suggested Learning Plan")
        plan = self.learning_path_gen.generate_path(result.missing_skills, result.resume_skills)

        # tabular view
        rows = []
        for p in plan:
            rows.append({'skill': p['skill'], 'time_estimate': p.get('time_estimate', ''), 'priority': p.get('priority', ''), 'resources': ', '.join(p.get('resources', []))})

        dfp = pd.DataFrame(rows)
        st.dataframe(dfp, use_container_width=True)

        # Timeline
        st.subheader('Estimated Timeline')
        timeline_data = []
        week_pattern = re.compile(r"(\d+)(?:\s*-\s*(\d+))?")
        now = datetime.now()
        for item in plan:
            time_str = str(item.get('time_estimate', '')).strip()
            weeks = None
            m = week_pattern.search(time_str)
            if m:
                try:
                    start_w = int(m.group(1))
                    end_w = int(m.group(2)) if m.group(2) else None
                    # prefer end_w as the total duration when provided, otherwise use start_w
                    weeks = end_w if end_w is not None else start_w
                except Exception:
                    weeks = None
            if weeks is None:
                # default estimate when not parseable (e.g., 'Varies')
                weeks = 8

            start_date = now
            finish_date = start_date + timedelta(weeks=weeks)
            timeline_data.append({'Skill': item.get('skill', 'Unknown'), 'Start': start_date, 'Finish': finish_date, 'Priority': item.get('priority', 'MEDIUM')})

        dft = pd.DataFrame(timeline_data)
        if not dft.empty:
            # ensure Start/Finish are datetime
            dft['Start'] = pd.to_datetime(dft['Start'])
            dft['Finish'] = pd.to_datetime(dft['Finish'])
            fig = px.timeline(dft, x_start='Start', x_end='Finish', y='Skill', color='Priority', title='Learning Timeline')
            fig.update_layout(yaxis={'autorange': 'reversed'})
            st.plotly_chart(fig, use_container_width=True)

        # Export plan
        st.subheader('Export Learning Plan')
        csv_plan = dfp.to_csv(index=False)
        st.download_button('Download CSV', csv_plan, 'learning_plan.csv', 'text/csv')

def main():
    try:
        app = SkillGapApp()
        app.run()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.exception(e)

if __name__ == "__main__":
    main()








































































































































# # ==========================================================
# # 🎯 AI Skill Gap Analyzer — Premium Visual Version (Final)
# # ==========================================================

# import streamlit as st
# import pandas as pd
# import numpy as np
# import time
# import json
# import plotly.express as px
# from sentence_transformers import SentenceTransformer
# from sklearn.metrics.pairwise import cosine_similarity

# # ==================== PAGE CONFIG ====================
# st.set_page_config(
#     page_title="AI Skill Gap Analyzer",
#     page_icon="🧠",
#     layout="wide",
#     initial_sidebar_state="expanded"
# )

# # ==================== CUSTOM STYLING ====================
# st.markdown("""
#     <style>
#         /* Background Gradient */
#         .stApp {
#             background: linear-gradient(to right, #e8f5e9, #e0f7fa);
#         }

#         /* Title Styling */
#         h1, h2, h3 {
#             color: #004d40 !important;
#             text-shadow: 1px 1px 2px #b2dfdb;
#         }

#         /* Buttons */
#         .stButton button {
#             background-color: #26a69a;
#             color: white;
#             border-radius: 12px;
#             border: none;
#             padding: 8px 20px;
#             font-size: 16px;
#             transition: 0.3s;
#         }
#         .stButton button:hover {
#             background-color: #00796b;
#             transform: scale(1.05);
#         }

#         /* Upload Box */
#         .upload-box {
#             border: 2px dashed #80cbc4;
#             padding: 25px;
#             border-radius: 15px;
#             background-color: #f9fbe7;
#         }

#         /* Skill Tag */
#         .skill-tag {
#             display: inline-block;
#             background-color: #81c784;
#             color: white;
#             padding: 6px 12px;
#             margin: 4px;
#             border-radius: 15px;
#             font-size: 14px;
#             box-shadow: 1px 1px 4px rgba(0,0,0,0.2);
#         }

#         /* Footer */
#         footer {
#             visibility: hidden;
#         }
#     </style>
# """, unsafe_allow_html=True)

# # ==================== HEADER ====================
# st.markdown("""
# <div style='text-align: center;'>
#     <img src="https://cdn-icons-png.flaticon.com/512/2721/2721298.png" width="100">
#     <h1 style='color: #00695c;'>AI Skill Gap Analyzer</h1>
#     <p style='font-size:18px;'>Empowering Careers through AI-driven Skill Insights 💼🤖</p>
# </div>
# """, unsafe_allow_html=True)

# # ==================== MODEL INITIALIZATION ====================
# @st.cache_resource
# def load_model():
#     return SentenceTransformer('paraphrase-MiniLM-L6-v2')

# model = load_model()

# # ==================== HELPER FUNCTIONS ====================
# def extract_skills(text):
#     text = text.lower()
#     skill_keywords = [
#         'python', 'java', 'machine learning', 'data analysis',
#         'cloud', 'docker', 'git', 'sql', 'tensorflow', 'pandas',
#         'communication', 'leadership', 'teamwork', 'creativity'
#     ]
#     found = [s for s in skill_keywords if s in text]
#     return list(set(found))

# def analyze_skill_gap(resume_text, jd_text):
#     resume_skills = extract_skills(resume_text)
#     jd_skills = extract_skills(jd_text)

#     resume_vecs = model.encode(resume_skills)
#     jd_vecs = model.encode(jd_skills)

#     sim = cosine_similarity(resume_vecs, jd_vecs)
#     matched, partial, missing = [], [], []

#     for i, rs in enumerate(resume_skills):
#         scores = sim[i]
#         if any(s >= 0.8 for s in scores):
#             matched.append(rs)
#         elif any(0.5 <= s < 0.8 for s in scores):
#             partial.append(rs)
#         else:
#             missing.append(rs)

#     for js in jd_skills:
#         if js not in resume_skills:
#             missing.append(js)

#     return {
#         'matched': list(set(matched)),
#         'partial': list(set(partial)),
#         'missing': list(set(missing)),
#         'resume_skills': resume_skills,
#         'jd_skills': jd_skills
#     }

# # ==================== STREAMLIT TABS ====================
# tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
#     "📄 Upload Documents",
#     "🧠 Extract Skills",
#     "🔍 Analysis",
#     "📊 Visualization",
#     "📈 Score & Report",
#     "🎯 Learning Path"
# ])

# # --------------------------------------------------------
# # TAB 1: Upload Documents
# # --------------------------------------------------------
# # --------------------------------------------------------
# # TAB 1: Upload Documents
# # --------------------------------------------------------
# with tab1:
#     st.image("https://cdn-icons-png.flaticon.com/512/1077/1077012.png", width=100)
#     st.subheader("Upload Your Resume & Job Description")

#     with st.container():
#         st.markdown('<div class="upload-box">', unsafe_allow_html=True)
#         st.write("### 📤 Upload Files or Paste Text Below")

#         col1, col2 = st.columns(2)
#         with col1:
#             uploaded_resume = st.file_uploader("📁 Upload Resume", type=["pdf", "docx", "txt"])
#         with col2:
#             uploaded_jd = st.file_uploader("📋 Upload Job Description", type=["pdf", "docx", "txt"])

#         # --- Helper functions for reading files ---
#         import PyPDF2
#         import docx

#         def read_pdf(file):
#             pdf_reader = PyPDF2.PdfReader(file)
#             return "\n".join([page.extract_text() or "" for page in pdf_reader.pages])

#         def read_docx(file):
#             doc = docx.Document(file)
#             return "\n".join([para.text for para in doc.paragraphs])

#         # --- Auto-fill text areas if files uploaded ---
#         resume_text = ""
#         jd_text = ""

#         if uploaded_resume:
#             if uploaded_resume.name.endswith(".pdf"):
#                 resume_text = read_pdf(uploaded_resume)
#             elif uploaded_resume.name.endswith(".docx"):
#                 resume_text = read_docx(uploaded_resume)
#             else:
#                 resume_text = uploaded_resume.read().decode("utf-8")
#             st.success(f"✅ Resume uploaded: {uploaded_resume.name}")

#         if uploaded_jd:
#             if uploaded_jd.name.endswith(".pdf"):
#                 jd_text = read_pdf(uploaded_jd)
#             elif uploaded_jd.name.endswith(".docx"):
#                 jd_text = read_docx(uploaded_jd)
#             else:
#                 jd_text = uploaded_jd.read().decode("utf-8")
#             st.success(f"✅ Job Description uploaded: {uploaded_jd.name}")

#         # --- Text Areas (editable even after upload) ---
#         resume_text = st.text_area("📄 Paste Resume Text", value=resume_text, height=150)
#         jd_text = st.text_area("📜 Paste Job Description Text", value=jd_text, height=150)
#         st.markdown('</div>', unsafe_allow_html=True)

#         # --- Analyze Button ---
#         if st.button("🔍 Analyze Skill Gap"):
#             if not resume_text.strip() or not jd_text.strip():
#                 st.warning("⚠️ Please upload or paste both Resume and Job Description.")
#             else:
#                 with st.spinner("Analyzing skills... please wait ⏳"):
#                     time.sleep(2)
#                     result = analyze_skill_gap(resume_text, jd_text)
#                     st.session_state['result'] = result
#                 st.success("✅ Analysis complete! Proceed to next tab.")

# # --------------------------------------------------------
# # TAB 2: Extract Skills
# # --------------------------------------------------------
# with tab2:
#     st.image("https://cdn-icons-png.flaticon.com/512/1903/1903162.png", width=100)
#     st.subheader("Extracted Skills")

#     if 'result' in st.session_state:
#         res = st.session_state['result']
#         st.write("### 🧾 Resume Skills:")
#         for s in res['resume_skills']:
#             st.markdown(f"<span class='skill-tag'>{s}</span>", unsafe_allow_html=True)

#         st.write("### 💼 Job Description Skills:")
#         for s in res['jd_skills']:
#             st.markdown(f"<span class='skill-tag' style='background-color:#64b5f6;'>{s}</span>", unsafe_allow_html=True)
#     else:
#         st.info("Please upload documents and run the analysis first.")

# # --------------------------------------------------------
# # TAB 3: Analysis
# # --------------------------------------------------------
# with tab3:
#     st.image("https://cdn-icons-png.flaticon.com/512/4149/4149670.png", width=100)
#     st.subheader("Skill Gap Analysis")

#     if 'result' in st.session_state:
#         res = st.session_state['result']
#         st.write("### ✅ Matched Skills:")
#         st.write(", ".join(res['matched']) or "None")

#         st.write("### ⚠️ Partial Matches:")
#         st.write(", ".join(res['partial']) or "None")

#         st.write("### ❌ Missing Skills:")
#         st.write(", ".join(res['missing']) or "None")
#     else:
#         st.info("Run the analysis in the first tab.")

# # --------------------------------------------------------
# # TAB 4: Visualization
# # --------------------------------------------------------
# with tab4:
#     if 'result' in st.session_state:
#         res = st.session_state['result']
#         data = {
#             'Category': ['Matched', 'Partial', 'Missing'],
#             'Count': [len(res['matched']), len(res['partial']), len(res['missing'])]
#         }
#         df = pd.DataFrame(data)
#         fig = px.pie(df, names='Category', values='Count',
#                      color_discrete_sequence=px.colors.sequential.Tealgrn)
#         st.plotly_chart(fig, use_container_width=True)
#     else:
#         st.info("Run the analysis first to view visualization.")

# # --------------------------------------------------------
# # TAB 5: Score & Report
# # --------------------------------------------------------
# with tab5:
#     if 'result' in st.session_state:
#         res = st.session_state['result']
#         score = round(len(res['matched']) / (len(res['jd_skills']) + 1e-5) * 100, 2)
#         st.metric(label="Overall Skill Match Score", value=f"{score}%")
#         st.progress(int(score))
#         st.download_button("📥 Download JSON Report",
#                            json.dumps(res, indent=2),
#                            file_name="skill_gap_report.json",
#                            mime="application/json")
#     else:
#         st.info("Please complete the analysis first.")

# # --------------------------------------------------------
# # TAB 6: Learning Path
# # --------------------------------------------------------
# with tab6:
#     if 'result' in st.session_state:
#         res = st.session_state['result']
#         missing = res['missing']

#         st.image("https://cdn-icons-png.flaticon.com/512/4727/4727413.png", width=100)
#         st.subheader("🎯 Personalized Learning Path")

#         if missing:
#             plan = []
#             for skill in missing:
#                 plan.append({
#                     "Skill": skill,
#                     "Priority": np.random.choice(["High", "Medium", "Low"]),
#                     "Estimated Time (hrs)": np.random.randint(8, 20),
#                     "Suggested Resource": f"https://www.coursera.org/search?query={skill}"
#                 })
#             df_plan = pd.DataFrame(plan)
#             st.dataframe(df_plan, use_container_width=True)

#             st.download_button("📥 Download Learning Plan (CSV)",
#                                df_plan.to_csv(index=False),
#                                file_name="learning_path.csv",
#                                mime="text/csv")
#         else:
#             st.success("🎉 Great! No missing skills detected.")
#     else:
#         st.info("Run the skill gap analysis first.")

# # --------------------------------------------------------
# # FOOTER
# # --------------------------------------------------------
# st.markdown("""
# <hr style='border: 1px solid #80cbc4;'>
# <div style='text-align:center; color: #004d40;'>
#     <i>“Skill is the unified force of experience, intellect and passion.”</i><br>
#     <b>- John Ruskin</b>
# </div>
# """, unsafe_allow_html=True)
