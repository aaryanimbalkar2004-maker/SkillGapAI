# """
# Complete Document Processing Pipeline for AI Skill Gap Analyzer
# Milestone 1: Data Ingestion and Parsing

# Requirements:
# pip install streamlit PyPDF2 python-docx textract pandas scikit-learn matplotlib
# Run with: streamlit run complete_pipeline.py
# """

import streamlit as st
import PyPDF2
import docx
import pandas as pd
import re
import logging
from typing import Dict, List, Tuple
from io import BytesIO

 # --- Extra Imports for Analysis ---
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt

 # ---------------- Logging ----------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

 # ---------------- Text Extractor ----------------
class TextExtractor:
    """Extracts text from PDF, DOCX, or TXT files."""
    def extract_pdf(self, file: BytesIO) -> str:
        text = ""
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def extract_docx(self, file: BytesIO) -> str:
        text = ""
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def extract_txt(self, file: BytesIO) -> str:
        return file.read().decode("utf-8")

    def extract(self, uploaded_file) -> str:
        filename = uploaded_file.name.lower()
        if filename.endswith(".pdf"):
            return self.extract_pdf(uploaded_file)
        elif filename.endswith(".docx"):
            return self.extract_docx(uploaded_file)
        elif filename.endswith(".txt"):
            return self.extract_txt(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {uploaded_file.name}")

# # ---------------- Text Cleaner ----------------
class TextCleaner:
    """Cleans extracted text by removing extra spaces and unwanted characters."""
    def clean(self, text: str) -> str:
        text = re.sub(r"\s+", " ", text)
        text = text.strip()
        return text

# ---------------- Document Parser ----------------
class DocumentParser:
    def parse_pdf(self, file: BytesIO) -> str:
        text = ""
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def parse_docx(self, file: BytesIO) -> str:
        text = ""
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def parse_txt(self, file: BytesIO) -> str:
        return file.read().decode("utf-8")

    def parse_file(self, uploaded_file) -> Tuple[str, str]:
        filename = uploaded_file.name.lower()
        if filename.endswith(".pdf"):
            return "pdf", self.parse_pdf(uploaded_file)
        elif filename.endswith(".docx"):
            return "docx", self.parse_docx(uploaded_file)
        elif filename.endswith(".txt"):
            return "txt", self.parse_txt(uploaded_file)
        else:
            raise ValueError("Unsupported file format")

# # ---------------- Document Processor ----------------
class DocumentProcessor:
    def __init__(self):
        self.parser = DocumentParser()
        self.processed_docs: List[Dict] = []

    def process_document(self, uploaded_file, doc_type: str) -> Dict:
        try:
            ext, text = self.parser.parse_file(uploaded_file)
            
            # Use TextCleaner to clean text
            cleaner = TextCleaner()
            clean_text = cleaner.clean(text)
            
            return {
                "filename": uploaded_file.name,
                "document_type": doc_type,
                "extension": ext,
                "content": clean_text,
                "success": True,
                "error": None,
            }
        except Exception as e:
            return {
                "filename": uploaded_file.name,
                "document_type": doc_type,
                "extension": None,
                "content": None,
                "success": False,
                "error": str(e),
            }

    def clean_text(self, text: str) -> str:
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def process_files(self, resumes, jds):
        self.processed_docs = []

        for file in resumes:
            self.processed_docs.append(self.process_document(file, "resume"))

        for file in jds:
            self.processed_docs.append(self.process_document(file, "job_description"))

        return self.processed_docs

    # ---------------- Display Processing Results ----------------
    def _display_processing_results(self):
        st.header("📊 Processing Results")

        if not self.processed_docs:
            st.info("No documents processed yet.")
            return

        for doc in self.processed_docs:
            with st.expander(f"📄 {doc['filename']} ({doc['document_type']})"):
                if doc["success"]:
                    st.success("✅ Processed Successfully")
                    st.write(doc["content"][:500] + "..." if len(doc["content"]) > 500 else doc["content"])
                else:
                    st.error(f"❌ Failed to process: {doc['error']}")

    # ---------------- Extra Analysis Features ----------------
    def _extra_analysis_features(self):
        

        processed_docs = self.processed_docs
        resumes = [doc for doc in processed_docs if doc["document_type"] == "resume" and doc["success"]]
        jds = [doc for doc in processed_docs if doc["document_type"] == "job_description" and doc["success"]]

        # --- Case 1: Resume + JD (ATS comparison)
        if resumes and jds:
            st.subheader("📊 ATS Analysis (Resume vs Job Description)")

            resume_text = " ".join([doc["content"] for doc in resumes])
            jd_text = " ".join([doc["content"] for doc in jds])

            # Vectorize
            vectorizer = TfidfVectorizer(stop_words="english")
            tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

            ats_score = round(similarity * 100, 2)
            st.success(f"✅ ATS Score: {ats_score}%")

            # Bar chart
            fig, ax = plt.subplots()
            ax.bar(["ATS Score"], [ats_score],
                   color="green" if ats_score >= 80 else "orange" if ats_score >= 60 else "red")
            ax.set_ylim(0, 100)
            ax.set_ylabel("Score (%)")
            st.pyplot(fig)

            # Keyword analysis
            resume_words = set(resume_text.lower().split())
            jd_words = set(jd_text.lower().split())
            common_words = resume_words.intersection(jd_words)
            missing_words = jd_words - resume_words

            st.info(f"🔍 Common Keywords: {len(common_words)} | ❌ Missing Keywords: {len(missing_words)}")
            if missing_words:
                st.write("**⚠️ Keywords missing in Resume (important for ATS):**")
                st.write(", ".join(list(missing_words)[:20]))

            # Suggestions
            st.subheader("💡 ATS Improvement Suggestions")
            suggestions = []
            if ats_score < 60:
                suggestions.append("Increase keyword overlap with the Job Description.")
            if len(resume_text.split()) < 250:
                suggestions.append("Resume content seems short. Add more details about your experience.")
            if "experience" not in resume_words:
                suggestions.append("Add a dedicated 'Experience' section.")
            if "skills" not in resume_words:
                suggestions.append("Add a clear 'Skills' section with keywords from the Job Description.")

            if suggestions:
                for s in suggestions:
                    st.write(f"- {s}")
            else:
                st.success("🎉 Your resume is well-optimized for this Job Description!")

        # --- Case 2: Resume only (Self-analysis)
        elif resumes and not jds:
            st.subheader("📑 Resume Analysis")

            resume_text = " ".join([doc["content"] for doc in resumes])
            word_count = len(resume_text.split())
            char_count = len(resume_text)

            st.info(f"📝 Word Count: {word_count}")
            st.info(f"🔠 Character Count: {char_count}")

            # Skill extraction (simple keyword-based)
            keywords = ["python", "java", "sql", "machine learning", "data", "cloud", "powerbi"]
            found = [kw for kw in keywords if kw.lower() in resume_text.lower()]
            st.success(f"✅ Skills Detected: {', '.join(found) if found else 'No major keywords found'}")

            # Section completeness check
            sections = ["summary", "experience", "projects", "skills", "education", "certifications"]
            missing_sections = [sec for sec in sections if sec not in resume_text.lower()]
            if missing_sections:
                st.warning(f"⚠️ Missing Sections: {', '.join(missing_sections)}")

            # Improvement tips
            if word_count < 250:
                st.warning("⚠️ Resume content seems short. Add more details about projects, experience, and achievements.")
            else:
                st.success("🎉 Resume length looks good!")

        else:
            st.warning("No documents available for analysis")

 #---------------- Main App ----------------#
def main():
    st.set_page_config(page_title="Documents Result: ", layout="wide")
    st.title("Documents Result")

    processor = DocumentProcessor()

    with st.sidebar:
        st.header("Upload Documents")
        resumes = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        jds = st.file_uploader("Upload Job Descriptions", type=["pdf", "docx", "txt"], accept_multiple_files=True)

        if st.button("Process Documents"):
            processed = processor.process_files(resumes, jds)
            if processed:
                st.session_state["processed_docs"] = processed
                st.success("Documents processed successfully!")
            else:
                st.error("No documents were processed.")

    if "processed_docs" in st.session_state:
        processor.processed_docs = st.session_state["processed_docs"]
        processor._display_processing_results()
        processor._extra_analysis_features()

if __name__ == "__main__":
    main()




